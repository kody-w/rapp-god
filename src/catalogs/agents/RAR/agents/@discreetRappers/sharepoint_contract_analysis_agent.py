"""
Agent: ContractAnalysisAgent
Purpose: Analyze, interpret, and summarize contracts stored in Azure File Storage
Data Sources: Azure File Storage (contracts/ folder), Azure OpenAI for analysis
Production Ready: Reads real documents, extracts text, performs AI-powered analysis

Supported formats: PDF, DOCX, TXT
Storage path: contracts/ folder in Azure File Storage
"""

from __future__ import annotations

# ═══════════════════════════════════════════════════════════════
# RAPP AGENT MANIFEST — Do not remove. Used by registry builder.
# ═══════════════════════════════════════════════════════════════
__manifest__ = {
    "schema": "rapp-agent/1.0",
    "name": "@discreetRappers/sharepoint_contract_analysis_agent",
    "version": "1.0.2",
    "display_name": "ContractAnalysis",
    "description": "Analyzes contract documents in Azure File Storage with Azure OpenAI \u2014 clause extraction, risk flagging, and comparison.",
    "author": "Bill Whalen",
    "tags": ["integrations", "sharepoint", "contracts", "analysis", "legal"],
    "category": "integrations",
    "quality_tier": "community",
    "requires_env": ["AZURE_FILES_SHARE_NAME", "AZURE_OPENAI_API_VERSION", "AZURE_OPENAI_DEPLOYMENT_NAME", "AZURE_OPENAI_ENDPOINT", "AZURE_STORAGE_ACCOUNT_NAME"],
    "dependencies": ["@rapp/basic_agent"],
}
# ═══════════════════════════════════════════════════════════════


import json
import logging
import os
import io
import re
from datetime import datetime
from typing import Optional, Dict, List, Any
from agents.basic_agent import BasicAgent

# Document processing imports
# Note: Auto-installation of missing packages is handled globally by function_app.py
# These imports will trigger auto-install if the packages are missing

# PDF support - try pypdf (modern) first, then PyPDF2 (legacy)
PDF_SUPPORT = False
pypdf_module = None

try:
    import pypdf
    pypdf_module = pypdf
    PDF_SUPPORT = True
except ImportError:
    try:
        import PyPDF2
        pypdf_module = PyPDF2
        PDF_SUPPORT = True
    except ImportError:
        logging.warning("PDF support disabled - pypdf/PyPDF2 not available")

# PDF generation support (reportlab)
PDF_GENERATION = False
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    PDF_GENERATION = True
except ImportError:
    logging.warning("PDF generation disabled - reportlab not available")

# DOCX support
DOCX_SUPPORT = False
DocxDocument = None

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    logging.warning("DOCX support disabled - python-docx not available")

# Azure imports
try:
    from azure.identity import DefaultAzureCredential
    from azure.storage.fileshare import ShareFileClient, ShareDirectoryClient, ShareServiceClient
    from openai import AzureOpenAI
    from azure.identity import get_bearer_token_provider
    AZURE_SUPPORT = True
except ImportError as e:
    AZURE_SUPPORT = False
    logging.warning(f"Azure SDK not fully installed: {e}")


class ContractAnalysisAgent(BasicAgent):
    """
    Production contract analysis agent that reads documents from Azure File Storage
    and uses Azure OpenAI to extract clauses, generate summaries, and identify risks.

    Storage Structure:
        contracts/           - Root folder for all contracts
        contracts/templates/ - Standard contract templates for comparison
        contracts/analysis/  - Stored analysis results (optional)

    Supported Actions:
        - list_contracts: List available contracts in storage
        - analyze_contract: Full analysis with all extractions
        - extract_clauses: Extract specific clause categories
        - summarize_contract: Generate executive summary
        - identify_risks: Compare against standard terms
        - compare_contracts: Compare two contracts side-by-side
    """

    def __init__(self):
        self.name = 'ContractAnalysis'
        self.metadata = {
            "name": self.name,
            "description": "Analyzes recording and entertainment contracts FROM THE LABEL'S PERSPECTIVE. Identifies risks to the label, flags artist-favorable terms, extracts clauses, and generates executive summaries for label decision-makers. Supports PDF, DOCX, and TXT formats.",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "description": "Action to perform. Use 'full_workup' when user says 'work on' a contract to run comprehensive analysis.",
                        "enum": [
                            "list_contracts",
                            "full_workup",
                            "analyze_contract",
                            "extract_clauses",
                            "summarize_contract",
                            "identify_risks",
                            "compare_contracts"
                        ]
                    },
                    "contract_name": {
                        "type": "string",
                        "description": "Name of the contract file in Azure storage (e.g., 'artist_agreement_2026.pdf')"
                    },
                    "contract_name_b": {
                        "type": "string",
                        "description": "Second contract name for comparison (used with compare_contracts action)"
                    },
                    "clause_types": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Specific clause types to extract: financial, rights, obligations, termination, exclusivity, territory, duration"
                    },
                    "summary_type": {
                        "type": "string",
                        "description": "Type of summary: executive (brief), detailed, or legal",
                        "enum": ["executive", "detailed", "legal"]
                    },
                    "audience": {
                        "type": "string",
                        "description": "Target audience for summary: legal, business, executive",
                        "enum": ["legal", "business", "executive"]
                    }
                },
                "required": ["action"]
            }
        }
        super().__init__(name=self.name, metadata=self.metadata)

        # Initialize Azure clients
        self.storage_account = os.environ.get('AZURE_STORAGE_ACCOUNT_NAME', 'stov4bzgynnlvii')
        self.share_name = os.environ.get('AZURE_FILES_SHARE_NAME', 'azfrapp-ov4bzgynnlviiov4bzgynnlvii')
        self.contracts_folder = 'contracts'

        # Initialize OpenAI client
        self.openai_client = None
        self.deployment_name = os.environ.get('AZURE_OPENAI_DEPLOYMENT_NAME', 'gpt-5.1-chat')
        self._init_openai_client()

    def _init_openai_client(self):
        """Initialize Azure OpenAI client with managed identity."""
        try:
            endpoint = os.environ.get('AZURE_OPENAI_ENDPOINT')
            if not endpoint:
                logging.warning("AZURE_OPENAI_ENDPOINT not set")
                return

            credential = DefaultAzureCredential()
            token_provider = get_bearer_token_provider(
                credential,
                "https://cognitiveservices.azure.com/.default"
            )

            self.openai_client = AzureOpenAI(
                azure_endpoint=endpoint,
                azure_ad_token_provider=token_provider,
                api_version=os.environ.get('AZURE_OPENAI_API_VERSION', '2025-01-01-preview')
            )
            logging.info("ContractAnalysisAgent: OpenAI client initialized")
        except Exception as e:
            logging.error(f"Failed to initialize OpenAI client: {e}")

    def _get_share_service_client(self) -> Optional[ShareServiceClient]:
        """Get Azure File Share service client."""
        try:
            credential = DefaultAzureCredential()
            account_url = f"https://{self.storage_account}.file.core.windows.net"
            return ShareServiceClient(
                account_url=account_url,
                credential=credential,
                token_intent="backup"  # Required for token-based auth
            )
        except Exception as e:
            logging.error(f"Failed to create share service client: {e}")
            return None

    def _list_files_in_folder(self, folder_path: str = None) -> List[Dict]:
        """List all files in the contracts folder."""
        try:
            service_client = self._get_share_service_client()
            if not service_client:
                return []

            share_client = service_client.get_share_client(self.share_name)
            target_folder = folder_path or self.contracts_folder

            try:
                directory_client = share_client.get_directory_client(target_folder)
                files = []

                for item in directory_client.list_directories_and_files():
                    if not item.get('is_directory', False):
                        file_name = item['name']
                        # Get file properties
                        file_client = directory_client.get_file_client(file_name)
                        props = file_client.get_file_properties()

                        files.append({
                            "name": file_name,
                            "size_kb": round(props.size / 1024, 2),
                            "last_modified": props.last_modified.isoformat() if props.last_modified else None,
                            "path": f"{target_folder}/{file_name}"
                        })

                return files
            except Exception as e:
                if "ResourceNotFound" in str(e):
                    logging.info(f"Folder {target_folder} does not exist, will be created on first upload")
                    return []
                raise

        except Exception as e:
            logging.error(f"Error listing files: {e}")
            return []

    def _read_file_content(self, file_path: str) -> Optional[bytes]:
        """Read file content from Azure File Storage."""
        try:
            service_client = self._get_share_service_client()
            if not service_client:
                return None

            share_client = service_client.get_share_client(self.share_name)
            file_client = share_client.get_file_client(file_path)

            download = file_client.download_file()
            return download.readall()

        except Exception as e:
            logging.error(f"Error reading file {file_path}: {e}")
            return None

    def _write_file_content(self, file_path: str, content: str) -> bool:
        """Write content to Azure File Storage."""
        try:
            service_client = self._get_share_service_client()
            if not service_client:
                return False

            share_client = service_client.get_share_client(self.share_name)

            # Ensure directory exists
            dir_path = '/'.join(file_path.split('/')[:-1])
            if dir_path:
                try:
                    dir_client = share_client.get_directory_client(dir_path)
                    dir_client.create_directory()
                except Exception:
                    pass  # Directory may already exist

            file_client = share_client.get_file_client(file_path)
            content_bytes = content.encode('utf-8')
            file_client.upload_file(content_bytes)

            logging.info(f"Successfully wrote file: {file_path}")
            return True

        except Exception as e:
            logging.error(f"Error writing file {file_path}: {e}")
            return False

    def _generate_download_url(self, file_path: str) -> str:
        """Generate a download URL for the file.

        Note: This storage account uses Entra ID authentication only (shared key access disabled).
        The returned URL requires authentication to access. Users can:
        1. Open in Azure Portal to download
        2. Use Azure Storage Explorer with their credentials
        3. Access via authenticated API calls
        """
        account_url = f"https://{self.storage_account}.file.core.windows.net"
        file_url = f"{account_url}/{self.share_name}/{file_path}"
        return file_url

    def _save_analysis_report(self, contract_name: str, analysis_data: Dict) -> Dict:
        """Save analysis report as professional PDF and return download info."""
        try:
            # Generate report filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_name = contract_name.rsplit('.', 1)[0]  # Remove extension
            report_name = f"{base_name}_analysis_{timestamp}.pdf"
            report_path = f"contracts/analysis/{report_name}"

            # Generate PDF report
            if PDF_GENERATION:
                pdf_bytes = self._generate_pdf_report(contract_name, analysis_data)
                if pdf_bytes:
                    # Write PDF to storage
                    if self._write_file_bytes(report_path, pdf_bytes):
                        download_url = self._generate_download_url(report_path)
                        return {
                            "saved": True,
                            "format": "PDF",
                            "report_name": report_name,
                            "report_path": report_path,
                            "download_url": download_url,
                            "size_kb": round(len(pdf_bytes) / 1024, 2)
                        }

            # Fallback to JSON if PDF generation fails
            report_name = f"{base_name}_analysis_{timestamp}.json"
            report_path = f"contracts/analysis/{report_name}"
            report_content = json.dumps(analysis_data, indent=2, default=str)

            if self._write_file_content(report_path, report_content):
                download_url = self._generate_download_url(report_path)
                return {
                    "saved": True,
                    "format": "JSON",
                    "report_name": report_name,
                    "report_path": report_path,
                    "download_url": download_url,
                    "size_kb": round(len(report_content) / 1024, 2)
                }
            else:
                return {"saved": False, "error": "Failed to write file"}

        except Exception as e:
            logging.error(f"Error saving analysis report: {e}")
            return {"saved": False, "error": str(e)}

    def _write_file_bytes(self, file_path: str, content: bytes) -> bool:
        """Write binary content to Azure File Storage."""
        try:
            service_client = self._get_share_service_client()
            if not service_client:
                return False

            share_client = service_client.get_share_client(self.share_name)

            # Ensure directory exists
            dir_path = '/'.join(file_path.split('/')[:-1])
            if dir_path:
                try:
                    dir_client = share_client.get_directory_client(dir_path)
                    dir_client.create_directory()
                except Exception:
                    pass  # Directory may already exist

            file_client = share_client.get_file_client(file_path)
            file_client.upload_file(content)

            logging.info(f"Successfully wrote binary file: {file_path}")
            return True

        except Exception as e:
            logging.error(f"Error writing binary file {file_path}: {e}")
            return False

    def _generate_pdf_report(self, contract_name: str, analysis_data: Dict) -> Optional[bytes]:
        """Generate a professional PDF analysis report."""
        if not PDF_GENERATION:
            return None

        try:
            # === DEBUG LOGGING: Dump structure of all data sections ===
            logging.info("=" * 60)
            logging.info("PDF GENERATION - DATA STRUCTURE ANALYSIS")
            logging.info("=" * 60)

            # Check executive_summary
            exec_summary = analysis_data.get('executive_summary', {})
            if isinstance(exec_summary, dict):
                logging.info(f"executive_summary keys: {list(exec_summary.keys())}")
                if exec_summary.get('parse_error'):
                    logging.error(f"executive_summary has PARSE ERROR")
                logging.info(f"  - summary length: {len(exec_summary.get('summary', ''))}")
                logging.info(f"  - risk_level: {exec_summary.get('risk_level', 'MISSING')}")
                logging.info(f"  - key_points count: {len(exec_summary.get('key_points', []))}")
            else:
                logging.error(f"executive_summary is NOT a dict: {type(exec_summary)}")

            # Check risk_assessment
            risk_assessment = analysis_data.get('risk_assessment', {})
            if isinstance(risk_assessment, dict):
                logging.info(f"risk_assessment keys: {list(risk_assessment.keys())}")
                if risk_assessment.get('parse_error'):
                    logging.error(f"risk_assessment has PARSE ERROR - raw: {risk_assessment.get('raw_analysis', '')[:500]}")
                logging.info(f"  - overall_risk_level: {risk_assessment.get('overall_risk_level', 'MISSING')}")
                logging.info(f"  - risk_score: {risk_assessment.get('risk_score', 'MISSING')}")
                logging.info(f"  - risks count: {len(risk_assessment.get('risks', []))}")
                logging.info(f"  - summary length: {len(risk_assessment.get('summary', ''))}")
            else:
                logging.error(f"risk_assessment is NOT a dict: {type(risk_assessment)}")

            # Check full_analysis
            full_analysis = analysis_data.get('full_analysis', {})
            if isinstance(full_analysis, dict):
                logging.info(f"full_analysis keys: {list(full_analysis.keys())}")
                if full_analysis.get('parse_error'):
                    logging.error(f"full_analysis has PARSE ERROR - raw: {full_analysis.get('raw_analysis', '')[:500]}")
                logging.info(f"  - contract_type: {full_analysis.get('contract_type', 'MISSING')}")
                logging.info(f"  - parties count: {len(full_analysis.get('parties', []))}")
                logging.info(f"  - financial_terms keys: {list(full_analysis.get('financial_terms', {}).keys()) if isinstance(full_analysis.get('financial_terms'), dict) else 'NOT DICT'}")
            else:
                logging.error(f"full_analysis is NOT a dict: {type(full_analysis)}")

            # Check extracted_clauses
            extracted_clauses = analysis_data.get('extracted_clauses', {})
            if isinstance(extracted_clauses, dict):
                logging.info(f"extracted_clauses keys: {list(extracted_clauses.keys())}")
                if extracted_clauses.get('parse_error'):
                    logging.error(f"extracted_clauses has PARSE ERROR")

            logging.info("=" * 60)

            # Create PDF in memory
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                leftMargin=0.75*inch,
                rightMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )

            # Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                spaceAfter=20,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#1a365d')
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceBefore=20,
                spaceAfter=10,
                textColor=colors.HexColor('#2c5282')
            )
            subheading_style = ParagraphStyle(
                'CustomSubHeading',
                parent=styles['Heading3'],
                fontSize=12,
                spaceBefore=12,
                spaceAfter=6,
                textColor=colors.HexColor('#4a5568')
            )
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=8,
                leading=14,
                alignment=TA_JUSTIFY
            )
            bullet_style = ParagraphStyle(
                'CustomBullet',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                leftIndent=20,
                bulletIndent=10
            )
            risk_high = ParagraphStyle(
                'RiskHigh',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#c53030'),
                spaceAfter=8
            )
            risk_medium = ParagraphStyle(
                'RiskMedium',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#dd6b20'),
                spaceAfter=8
            )
            risk_low = ParagraphStyle(
                'RiskLow',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#38a169'),
                spaceAfter=8
            )

            story = []

            # Title
            story.append(Paragraph("CONTRACT ANALYSIS REPORT", title_style))
            story.append(Spacer(1, 0.1*inch))

            # Contract info header
            story.append(Paragraph(f"<b>Contract:</b> {contract_name}", body_style))
            story.append(Paragraph(f"<b>Analysis Date:</b> {analysis_data.get('analyzed_at', datetime.now().isoformat())}", body_style))
            story.append(Spacer(1, 0.2*inch))

            # Horizontal line
            story.append(Table([['']], colWidths=[7*inch], rowHeights=[2]))
            story[-1].setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#2c5282'))]))
            story.append(Spacer(1, 0.2*inch))

            # Executive Summary Section
            exec_summary = analysis_data.get('executive_summary', {})
            if exec_summary:
                story.append(Paragraph("EXECUTIVE SUMMARY", heading_style))

                if isinstance(exec_summary, dict):
                    summary_text = exec_summary.get('summary', '')
                    if summary_text:
                        # Use larger max_length for executive summary - don't truncate
                        story.append(Paragraph(self._clean_text(summary_text, max_length=2000), body_style))

                    # Risk Level Box
                    risk_level = exec_summary.get('risk_level', 'unknown').upper()
                    risk_style = risk_high if risk_level == 'HIGH' else (risk_medium if risk_level == 'MEDIUM' else risk_low)
                    story.append(Spacer(1, 0.1*inch))
                    story.append(Paragraph(f"<b>Overall Risk Level: {risk_level}</b>", risk_style))

                    # Key Points
                    key_points = exec_summary.get('key_points', [])
                    if key_points:
                        story.append(Paragraph("<b>Key Points:</b>", subheading_style))
                        for point in key_points[:10]:  # Limit to 10 points
                            if isinstance(point, dict):
                                # Format dict with point and clickable ref
                                point_text = point.get('point', '')
                                point_ref = point.get('ref', '')
                                clickable_ref = self._format_clickable_ref(point_ref, analysis_data.get('_contract_text', '')) if point_ref else ""
                                formatted = f"{self._clean_text(point_text, max_length=300)} {clickable_ref}"
                            else:
                                formatted = self._clean_text(str(point), max_length=300)
                            story.append(Paragraph(f"* {formatted}", bullet_style))

                    # Recommendation
                    recommendation = exec_summary.get('recommendation', '')
                    if recommendation:
                        story.append(Spacer(1, 0.1*inch))
                        story.append(Paragraph(f"<b>Recommendation:</b> {self._clean_text(recommendation, max_length=1000)}", body_style))

            story.append(PageBreak())

            # Risk Assessment Section
            risk_assessment = analysis_data.get('risk_assessment', {})
            story.append(Paragraph("RISK ASSESSMENT", heading_style))

            if isinstance(risk_assessment, dict):
                # Check for parse error
                if risk_assessment.get('parse_error'):
                    story.append(Paragraph(
                        "<b><font color='red'>Warning: Risk assessment data extraction encountered issues.</font></b>",
                        body_style
                    ))
                    # Show raw analysis if available
                    raw_analysis = risk_assessment.get('raw_analysis', '')
                    if raw_analysis and len(raw_analysis) > 50:
                        story.append(Paragraph("<b>Raw Analysis Output:</b>", subheading_style))
                        # Show first 2000 chars of raw analysis
                        story.append(Paragraph(self._clean_text(raw_analysis[:2000], max_length=2000), body_style))
                else:
                    # Overall risk
                    overall_risk = risk_assessment.get('overall_risk_level', 'Unknown').upper()
                    risk_score = risk_assessment.get('risk_score', 'N/A')
                    risk_style = risk_high if overall_risk == 'HIGH' else (risk_medium if overall_risk == 'MEDIUM' else risk_low)
                    story.append(Paragraph(f"<b>Risk Level: {overall_risk} | Score: {risk_score}/100</b>", risk_style))

                    # Summary
                    risk_summary = risk_assessment.get('summary', '')
                    if risk_summary:
                        story.append(Paragraph(self._clean_text(risk_summary, max_length=1000), body_style))
                    else:
                        story.append(Paragraph("<i>No risk summary available.</i>", body_style))

                    # Individual risks
                    risks = risk_assessment.get('risks', [])
                    if risks:
                        story.append(Paragraph("<b>Identified Risks:</b>", subheading_style))

                        # Create table cell style for wrapping text
                        cell_style = ParagraphStyle(
                            'CellStyle',
                            parent=styles['Normal'],
                            fontSize=8,
                            leading=10,
                            spaceAfter=0
                        )
                        header_cell_style = ParagraphStyle(
                            'HeaderCellStyle',
                            parent=styles['Normal'],
                            fontSize=9,
                            leading=11,
                            textColor=colors.white,
                            fontName='Helvetica-Bold'
                        )

                        # Build header row with Paragraphs
                        risk_data = [[
                            Paragraph('Category', header_cell_style),
                            Paragraph('Severity', header_cell_style),
                            Paragraph('Description', header_cell_style)
                        ]]

                        for risk in risks[:10]:  # Limit to 10 risks
                            if isinstance(risk, dict):
                                desc_text = self._clean_text(risk.get('description', 'N/A'), max_length=300)
                                risk_data.append([
                                    Paragraph(risk.get('category', 'N/A'), cell_style),
                                    Paragraph(risk.get('severity', 'N/A').upper(), cell_style),
                                    Paragraph(desc_text, cell_style)
                                ])

                        if len(risk_data) > 1:
                            # Adjusted widths: Category 1", Severity 0.7", Description 5"
                            risk_table = Table(risk_data, colWidths=[1*inch, 0.7*inch, 5*inch])
                            risk_table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
                                ('FONTSIZE', (0, 0), (-1, -1), 8),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
                                ('TOPPADDING', (0, 0), (-1, -1), 4),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                            ]))
                            story.append(risk_table)
                    else:
                        story.append(Paragraph("<i>No individual risks identified.</i>", body_style))

                    # Artist-favorable terms from risk assessment
                    artist_favorable = risk_assessment.get('artist_favorable_terms', [])
                    if artist_favorable:
                        story.append(Spacer(1, 0.15*inch))
                        story.append(Paragraph("<b>Artist-Favorable Terms (Label Concerns):</b>", subheading_style))
                        for term in artist_favorable[:6]:
                            if isinstance(term, dict):
                                term_text = term.get('term', 'N/A')
                                label_impact = term.get('label_impact', '')
                                ref = term.get('ref', '')
                                impact_text = f" - {label_impact}" if label_impact else ""
                                ref_text = f" (Ref: {ref})" if ref else ""
                                story.append(Paragraph(f"* {self._clean_text(str(term_text))}{impact_text}{ref_text}", bullet_style))

                    # Negotiation points
                    negotiation = risk_assessment.get('negotiation_points', risk_assessment.get('negotiation_priorities', []))
                    if negotiation:
                        story.append(Spacer(1, 0.15*inch))
                        story.append(Paragraph("<b>Recommended Negotiation Points:</b>", subheading_style))
                        for point in negotiation[:8]:
                            if isinstance(point, dict):
                                priority_text = point.get('priority', point.get('point', str(point)))
                                ref = point.get('ref', '')
                                ref_text = f" (Ref: {ref})" if ref else ""
                                story.append(Paragraph(f"* {self._clean_text(str(priority_text))}{ref_text}", bullet_style))
                            else:
                                story.append(Paragraph(f"* {self._clean_text(str(point))}", bullet_style))

                    # Deal breakers
                    deal_breakers = risk_assessment.get('deal_breakers', [])
                    if deal_breakers:
                        story.append(Spacer(1, 0.15*inch))
                        story.append(Paragraph("<b>Potential Deal Breakers:</b>", subheading_style))
                        for item in deal_breakers[:5]:
                            if isinstance(item, dict):
                                issue = item.get('issue', 'N/A')
                                ref = item.get('ref', '')
                                ref_text = f" (Ref: {ref})" if ref else ""
                                story.append(Paragraph(f"* {self._clean_text(str(issue))}{ref_text}", bullet_style))
                            else:
                                story.append(Paragraph(f"* {self._clean_text(str(item))}", bullet_style))
            else:
                story.append(Paragraph("<i>Risk assessment data not available.</i>", body_style))

            story.append(PageBreak())

            # Full Analysis Section
            full_analysis = analysis_data.get('full_analysis', {})
            logging.info(f"PDF Generation - full_analysis keys: {list(full_analysis.keys()) if isinstance(full_analysis, dict) else 'NOT DICT'}")

            story.append(Paragraph("DETAILED CONTRACT ANALYSIS", heading_style))

            if isinstance(full_analysis, dict):
                # Check for parse error
                if full_analysis.get('parse_error'):
                    story.append(Paragraph(
                        "<b><font color='red'>Warning: Detailed analysis data extraction encountered issues.</font></b>",
                        body_style
                    ))
                    # Show raw analysis if available
                    raw_analysis = full_analysis.get('raw_analysis', '')
                    if raw_analysis and len(raw_analysis) > 50:
                        story.append(Paragraph("<b>Raw Analysis Output:</b>", subheading_style))
                        # Show first 3000 chars of raw analysis
                        story.append(Paragraph(self._clean_text(raw_analysis[:3000], max_length=3000), body_style))
                else:
                    # Contract Type
                    contract_type = full_analysis.get('contract_type', 'Not identified')
                    logging.info(f"PDF Generation - contract_type: {contract_type}")
                    story.append(Paragraph(f"<b>Contract Type:</b> {contract_type}", body_style))

                    # Get contract text for snippets
                    contract_text = analysis_data.get('_contract_text', '')
                    logging.info(f"PDF Generation - contract_text length: {len(contract_text) if contract_text else 0}")

                    # Parties
                    parties = full_analysis.get('parties', [])
                    logging.info(f"PDF Generation - parties count: {len(parties) if parties else 0}")
                    if parties:
                        story.append(Paragraph("<b>Parties:</b>", subheading_style))
                        for party in parties:
                            if isinstance(party, dict):
                                try:
                                    party_ref = party.get('ref', '')
                                    clickable_ref = self._format_clickable_ref(party_ref, contract_text) if party_ref else ""
                                    story.append(Paragraph(f"* {party.get('name', 'N/A')} - {party.get('role', 'N/A')} {clickable_ref}", bullet_style))
                                except Exception as e:
                                    logging.error(f"Error rendering party: {e}")
                                    story.append(Paragraph(f"* {party.get('name', 'N/A')} - {party.get('role', 'N/A')}", bullet_style))

                    # Term - handle both string and dict formats
                    term = full_analysis.get('term_duration', '')
                    if term:
                        if isinstance(term, dict):
                            term_val = term.get('value', 'N/A')
                            term_ref = term.get('ref', '')
                            clickable_ref = self._format_clickable_ref(term_ref, contract_text) if term_ref else ""
                            story.append(Paragraph(f"<b>Term:</b> {self._clean_text(str(term_val))} {clickable_ref}", body_style))
                        else:
                            story.append(Paragraph(f"<b>Term:</b> {self._clean_text(str(term))}", body_style))

                    # Effective Date
                    effective_date = full_analysis.get('effective_date', {})
                    if effective_date and isinstance(effective_date, dict):
                        date_val = effective_date.get('value', '')
                        if date_val:
                            date_ref = effective_date.get('ref', '')
                            clickable_ref = self._format_clickable_ref(date_ref, contract_text) if date_ref else ""
                            story.append(Paragraph(f"<b>Effective Date:</b> {self._clean_text(str(date_val))} {clickable_ref}", body_style))

                    # Financial Terms
                    financial = full_analysis.get('financial_terms', {})
                    if financial and isinstance(financial, dict):
                        story.append(Paragraph("<b>Financial Terms:</b>", subheading_style))
                        # Use the formatter for complex nested values
                        if financial.get('advances'):
                            formatted_advances = self._format_value_for_pdf(financial.get('advances'))
                            story.append(Paragraph(f"* Advances: {formatted_advances}", bullet_style))
                        if financial.get('royalty_rates'):
                            formatted_royalties = self._format_value_for_pdf(financial.get('royalty_rates'))
                            story.append(Paragraph(f"* Royalty Rates: {formatted_royalties}", bullet_style))
                        if financial.get('payment_schedule'):
                            formatted_schedule = self._format_value_for_pdf(financial.get('payment_schedule'))
                            story.append(Paragraph(f"* Payment Schedule: {formatted_schedule}", bullet_style))
                        if financial.get('label_investment'):
                            formatted_investment = self._format_value_for_pdf(financial.get('label_investment'))
                            story.append(Paragraph(f"* Label Investment: {formatted_investment}", bullet_style))
                        if financial.get('recoupment_terms'):
                            formatted_recoup = self._format_value_for_pdf(financial.get('recoupment_terms'))
                            story.append(Paragraph(f"* Recoupment: {formatted_recoup}", bullet_style))

                    # Rights Secured (AI returns 'rights_secured' not 'rights_granted')
                    rights = full_analysis.get('rights_secured', []) or full_analysis.get('rights_granted', [])
                    if rights:
                        story.append(Paragraph("<b>Rights Secured:</b>", subheading_style))
                        for right in rights[:8]:
                            if isinstance(right, dict):
                                exclusivity = "Exclusive" if right.get('exclusivity') else "Non-exclusive"
                                right_desc = right.get('right', '') or right.get('description', 'N/A')
                                scope = right.get('scope', '')
                                duration = right.get('duration', '')
                                ref = right.get('ref', '')

                                details = [f"{exclusivity}"]
                                if scope:
                                    details.append(f"Territory: {scope}")
                                if duration:
                                    details.append(f"Duration: {duration}")

                                clickable_ref = self._format_clickable_ref(ref, contract_text) if ref else ""
                                story.append(Paragraph(
                                    f"* {self._clean_text(str(right_desc))} ({', '.join(details)}) {clickable_ref}",
                                    bullet_style
                                ))

                    # Artist Obligations
                    artist_obligations = full_analysis.get('artist_obligations', [])
                    if artist_obligations:
                        story.append(Paragraph("<b>Artist Obligations:</b>", subheading_style))
                        for obligation in artist_obligations[:6]:
                            if isinstance(obligation, dict):
                                obl_text = obligation.get('obligation', 'N/A')
                                deadline = obligation.get('deadline', '')
                                ref = obligation.get('ref', '')

                                deadline_text = f" (Deadline: {deadline})" if deadline else ""
                                clickable_ref = self._format_clickable_ref(ref, contract_text) if ref else ""
                                story.append(Paragraph(
                                    f"* {self._clean_text(str(obl_text))}{deadline_text} {clickable_ref}",
                                    bullet_style
                                ))

                    # Label Obligations
                    label_obligations = full_analysis.get('label_obligations', [])
                    if label_obligations:
                        story.append(Paragraph("<b>Label Obligations:</b>", subheading_style))
                        for obligation in label_obligations[:6]:
                            if isinstance(obligation, dict):
                                obl_text = obligation.get('obligation', 'N/A')
                                impact = obligation.get('financial_impact', '')
                                ref = obligation.get('ref', '')

                                impact_text = f" (Cost: {impact})" if impact else ""
                                clickable_ref = self._format_clickable_ref(ref, contract_text) if ref else ""
                                story.append(Paragraph(
                                    f"* {self._clean_text(str(obl_text))}{impact_text} {clickable_ref}",
                                    bullet_style
                                ))

                    # Label Protections
                    protections = full_analysis.get('label_protections', [])
                    if protections:
                        story.append(Paragraph("<b>Label Protections:</b>", subheading_style))
                        for protection in protections[:6]:
                            if isinstance(protection, dict):
                                clause = protection.get('clause', 'N/A')
                                ref = protection.get('ref', '')
                                clickable_ref = self._format_clickable_ref(ref, contract_text) if ref else ""
                                story.append(Paragraph(f"* {self._clean_text(str(clause))} {clickable_ref}", bullet_style))

                    # Termination Clauses
                    termination = full_analysis.get('termination_clauses', [])
                    if termination:
                        story.append(Paragraph("<b>Termination Provisions:</b>", subheading_style))
                        for term_clause in termination[:6]:
                            if isinstance(term_clause, dict):
                                trigger = term_clause.get('trigger', 'N/A')
                                who = term_clause.get('who_can_trigger', '')
                                impact = term_clause.get('label_impact', '')
                                ref = term_clause.get('ref', '')

                                who_text = f" (By: {who})" if who else ""
                                clickable_ref = self._format_clickable_ref(ref, contract_text) if ref else ""
                                story.append(Paragraph(
                                    f"* {self._clean_text(str(trigger))}{who_text} {clickable_ref}",
                                    bullet_style
                                ))

                    # Artist-Favorable Terms (risks to label)
                    artist_favorable = full_analysis.get('artist_favorable_terms', [])
                    if artist_favorable:
                        story.append(Paragraph("<b>Artist-Favorable Terms (Label Concerns):</b>", subheading_style))
                        for term in artist_favorable[:5]:
                            if isinstance(term, dict):
                                term_desc = term.get('term', 'N/A')
                                concern = term.get('concern', '')
                                ref = term.get('ref', '')

                                concern_text = f" - {concern}" if concern else ""
                                clickable_ref = self._format_clickable_ref(ref, contract_text) if ref else ""
                                story.append(Paragraph(
                                    f"* {self._clean_text(str(term_desc))}{concern_text} {clickable_ref}",
                                    bullet_style
                                ))

                    # Overall Assessment
                    assessment = full_analysis.get('overall_assessment', '')
                    if assessment:
                        story.append(Spacer(1, 0.15*inch))
                        story.append(Paragraph("<b>Overall Assessment:</b>", subheading_style))
                        story.append(Paragraph(self._clean_text(str(assessment)), body_style))
            else:
                story.append(Paragraph("<i>Detailed analysis data not available or could not be parsed.</i>", body_style))

            # SOURCE CONTRACT SECTION - Appended with page anchors for clickable references
            contract_text = analysis_data.get('_contract_text', '')
            if contract_text:
                story.append(PageBreak())
                story.append(Paragraph("SOURCE CONTRACT", heading_style))
                story.append(Paragraph(
                    "<i>The original contract text is included below. Click any page reference in the analysis above to jump directly to that location.</i>",
                    body_style
                ))
                story.append(Spacer(1, 0.2*inch))

                # Style for contract text (smaller, monospace-like)
                contract_style = ParagraphStyle(
                    'ContractText',
                    parent=styles['Normal'],
                    fontSize=9,
                    spaceAfter=6,
                    leading=12,
                    leftIndent=10,
                    rightIndent=10
                )

                # Split contract by page markers and create anchors
                # Look for [PAGE N] markers in the text
                page_pattern = re.compile(r'\[PAGE\s*(\d+)\]', re.IGNORECASE)

                # Split text by page markers, keeping the markers
                parts = page_pattern.split(contract_text)

                if len(parts) > 1:
                    # We have page markers
                    current_page = None
                    for i, part in enumerate(parts):
                        if i % 2 == 1:  # This is a page number
                            current_page = part
                            # Create anchor for this page
                            anchor_name = f"contract_page_{current_page}"
                            story.append(Spacer(1, 0.15*inch))
                            # Page header with anchor
                            story.append(Paragraph(
                                f'<a name="{anchor_name}"/><b>--- PAGE {current_page} ---</b>',
                                subheading_style
                            ))
                        else:
                            # This is content
                            if part.strip():
                                # Clean and split into paragraphs
                                paragraphs = part.strip().split('\n\n')
                                for para in paragraphs:
                                    if para.strip():
                                        clean_para = self._clean_text(para.strip(), max_length=2000)
                                        if clean_para:
                                            story.append(Paragraph(clean_para, contract_style))
                else:
                    # No page markers - create anchors for every ~2000 chars as "pages"
                    chunk_size = 2000
                    chunks = [contract_text[i:i+chunk_size] for i in range(0, len(contract_text), chunk_size)]
                    for page_num, chunk in enumerate(chunks, 1):
                        anchor_name = f"contract_page_{page_num}"
                        story.append(Spacer(1, 0.15*inch))
                        story.append(Paragraph(
                            f'<a name="{anchor_name}"/><b>--- PAGE {page_num} ---</b>',
                            subheading_style
                        ))
                        clean_chunk = self._clean_text(chunk.strip(), max_length=2500)
                        if clean_chunk:
                            story.append(Paragraph(clean_chunk, contract_style))

            # Footer
            story.append(Spacer(1, 0.5*inch))
            story.append(Table([['']], colWidths=[7*inch], rowHeights=[2]))
            story[-1].setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#2c5282'))]))
            story.append(Spacer(1, 0.1*inch))

            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=8,
                textColor=colors.grey,
                alignment=TA_CENTER
            )
            story.append(Paragraph(
                f"Generated by ContractAnalysis Agent | {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
                footer_style
            ))
            story.append(Paragraph(
                "This analysis is for informational purposes only and does not constitute legal advice.",
                footer_style
            ))

            # Build PDF
            doc.build(story)
            pdf_bytes = buffer.getvalue()
            buffer.close()

            logging.info(f"Generated PDF report: {len(pdf_bytes)} bytes")
            return pdf_bytes

        except Exception as e:
            logging.error(f"Error generating PDF report: {e}")
            return None

    def _clean_text(self, text: str, max_length: int = 500) -> str:
        """Clean text for PDF rendering - escape special characters and normalize Unicode."""
        if not text:
            return ""
        text = str(text)

        # Normalize Unicode characters that don't render in standard PDF fonts
        unicode_replacements = {
            '\u2011': '-',   # Non-breaking hyphen → regular hyphen
            '\u2010': '-',   # Hyphen → regular hyphen
            '\u2012': '-',   # Figure dash → regular hyphen
            '\u2013': '-',   # En-dash → regular hyphen
            '\u2014': '-',   # Em-dash → regular hyphen
            '\u2015': '-',   # Horizontal bar → regular hyphen
            '\u2018': "'",   # Left single quote → apostrophe
            '\u2019': "'",   # Right single quote → apostrophe
            '\u201a': "'",   # Single low quote → apostrophe
            '\u201b': "'",   # Single high-reversed quote → apostrophe
            '\u201c': '"',   # Left double quote → regular quote
            '\u201d': '"',   # Right double quote → regular quote
            '\u201e': '"',   # Double low quote → regular quote
            '\u201f': '"',   # Double high-reversed quote → regular quote
            '\u2022': '*',   # Bullet → asterisk
            '\u2023': '>',   # Triangular bullet → greater than
            '\u2024': '.',   # One dot leader → period
            '\u2025': '..',  # Two dot leader → two periods
            '\u2026': '...', # Ellipsis → three periods
            '\u2027': '-',   # Hyphenation point → hyphen
            '\u2032': "'",   # Prime → apostrophe
            '\u2033': '"',   # Double prime → quote
            '\u2039': '<',   # Single left angle quote
            '\u203a': '>',   # Single right angle quote
            '\u00ab': '<<',  # Left double angle quote
            '\u00bb': '>>',  # Right double angle quote
            '\u00a0': ' ',   # Non-breaking space → regular space
            '\u200b': '',    # Zero-width space → remove
            '\u200c': '',    # Zero-width non-joiner → remove
            '\u200d': '',    # Zero-width joiner → remove
            '\ufeff': '',    # BOM → remove
            '\u00b7': '*',   # Middle dot → asterisk
            '\u2212': '-',   # Minus sign → hyphen
            '\u00d7': 'x',   # Multiplication sign → x
            '\u00f7': '/',   # Division sign → slash
            '\u2248': '~',   # Almost equal → tilde
            '\u2260': '!=',  # Not equal → !=
            '\u2264': '<=',  # Less than or equal
            '\u2265': '>=',  # Greater than or equal
            '\u00b0': ' deg', # Degree symbol
            '\u00a9': '(c)', # Copyright
            '\u00ae': '(R)', # Registered
            '\u2122': '(TM)', # Trademark
        }

        for unicode_char, replacement in unicode_replacements.items():
            text = text.replace(unicode_char, replacement)

        # Replace any remaining non-ASCII characters that might cause issues
        # Keep basic extended ASCII (accented letters) but remove other oddities
        cleaned_chars = []
        for char in text:
            if ord(char) < 128:  # Standard ASCII
                cleaned_chars.append(char)
            elif ord(char) < 256:  # Extended ASCII (accented chars) - keep these
                cleaned_chars.append(char)
            else:  # Other Unicode - replace with space or skip
                cleaned_chars.append(' ')
        text = ''.join(cleaned_chars)

        # Clean up multiple spaces
        while '  ' in text:
            text = text.replace('  ', ' ')

        # Replace problematic characters for reportlab XML
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = text.replace('\n', ' ')
        text = text.replace('\r', '')

        # Limit length to prevent overflow
        if max_length and len(text) > max_length:
            text = text[:max_length - 3] + "..."
        return text.strip()

    def _format_clickable_ref(self, ref_text: str, contract_text: str = '') -> str:
        """Format a reference string as a clickable internal PDF link with snippet.

        Converts refs like "Page 3, Section 4.1" to clickable links that jump
        to the corresponding page in the appended contract text.
        Includes a short snippet from the referenced page for context.
        """
        if not ref_text or ref_text == 'N/A':
            return ""

        # Extract page number from reference (e.g., "Page 3" or "Pages 3-4")
        page_match = re.search(r'[Pp]age[s]?\s*(\d+)', str(ref_text))
        if page_match:
            page_num = page_match.group(1)
            anchor_name = f"contract_page_{page_num}"
            clean_ref = self._clean_text(str(ref_text), max_length=100)

            # Extract a snippet from the referenced page if contract text is available
            snippet = ""
            if contract_text:
                # Find the page marker and extract text after it
                page_pattern = re.compile(rf'\[PAGE\s*{page_num}\](.*?)(?:\[PAGE\s*\d+\]|$)', re.IGNORECASE | re.DOTALL)
                match = page_pattern.search(contract_text)
                if match:
                    page_content = match.group(1).strip()
                    # Get first 80 chars as snippet, clean it up
                    if page_content:
                        snippet_text = page_content[:120].replace('\n', ' ').strip()
                        # Truncate at word boundary
                        if len(snippet_text) >= 80:
                            last_space = snippet_text[:80].rfind(' ')
                            if last_space > 40:
                                snippet_text = snippet_text[:last_space]
                        snippet = self._clean_text(snippet_text, max_length=80)

            # Build the clickable reference with optional snippet
            if snippet:
                return f'<a href="#{anchor_name}" color="blue"><i>(Ref: {clean_ref})</i></a> <font size="8" color="gray">["{snippet}..."]</font>'
            else:
                return f'<a href="#{anchor_name}" color="blue"><i>(Ref: {clean_ref})</i></a>'
        else:
            # No page number found, just return plain ref
            clean_ref = self._clean_text(str(ref_text), max_length=100)
            return f"<i>(Ref: {clean_ref})</i>"

    def _format_value_for_pdf(self, value: Any, indent: int = 0) -> str:
        """Format a value (potentially nested dict/list) into readable text for PDF."""
        if value is None:
            return "N/A"

        if isinstance(value, str):
            return self._clean_text(value, max_length=None)

        if isinstance(value, (int, float, bool)):
            return str(value)

        if isinstance(value, list):
            if not value:
                return "None"
            # For simple lists, join with commas
            if all(isinstance(item, str) for item in value):
                return ", ".join(str(item) for item in value[:5])  # Limit to 5 items
            # For complex lists, format each item
            parts = []
            for item in value[:5]:
                parts.append(self._format_value_for_pdf(item, indent + 1))
            return "; ".join(parts)

        if isinstance(value, dict):
            # Format dict as readable key-value pairs
            parts = []
            for k, v in value.items():
                # Clean up key name (replace underscores, capitalize)
                key_name = k.replace('_', ' ').title()
                formatted_value = self._format_value_for_pdf(v, indent + 1)
                if formatted_value and formatted_value != "N/A":
                    parts.append(f"{key_name}: {formatted_value}")
            return "; ".join(parts) if parts else "N/A"

        return str(value)

    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF content using pypdf or PyPDF2.

        Includes clear page markers for document reference tracking.
        """
        if not PDF_SUPPORT or pypdf_module is None:
            return "[ERROR: PDF library not available. Install with: pip install pypdf]"

        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = pypdf_module.PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    # Clear page markers for AI to reference
                    text_parts.append(f"[PAGE {page_num + 1}]\n{page_text}\n[END PAGE {page_num + 1}]")

            return "\n\n".join(text_parts)
        except Exception as e:
            logging.error(f"PDF extraction error: {e}")
            return f"[ERROR: Failed to extract PDF text: {e}]"

    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        if not DOCX_SUPPORT or DocxDocument is None:
            return "[ERROR: python-docx not available. Install with: pip install python-docx]"

        try:
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as e:
            logging.error(f"DOCX extraction error: {e}")
            return f"[ERROR: Failed to extract DOCX text: {e}]"

    def _extract_text(self, file_path: str, content: bytes) -> str:
        """Extract text based on file extension."""
        ext = file_path.lower().split('.')[-1]

        if ext == 'pdf':
            return self._extract_text_from_pdf(content)
        elif ext in ['docx', 'doc']:
            return self._extract_text_from_docx(content)
        elif ext == 'txt':
            return content.decode('utf-8', errors='ignore')
        else:
            return f"[ERROR: Unsupported file format: {ext}. Supported: pdf, docx, txt]"

    def _extract_json_from_response(self, response: str) -> Dict:
        """Extract JSON from AI response, handling various formats."""
        if not response:
            logging.error("_extract_json_from_response: Empty response received")
            return {"parse_error": True, "raw_analysis": response, "error_type": "empty_response"}

        # Clean up common issues in AI responses
        cleaned_response = response.strip()

        # Remove control characters that can break JSON parsing
        cleaned_response = re.sub(r'[\x00-\x1f\x7f-\x9f]', ' ', cleaned_response)

        # Check for truncated JSON (common when hitting token limits)
        # If it starts with { but doesn't end with }, try to repair it
        if cleaned_response.startswith('{') and not cleaned_response.rstrip().endswith('}'):
            logging.warning("_extract_json_from_response: Detected possibly truncated JSON, attempting repair...")
            # Try to find the last complete key-value pair and close the JSON
            repaired = self._repair_truncated_json(cleaned_response)
            if repaired:
                try:
                    result = json.loads(repaired)
                    logging.info(f"_extract_json_from_response: Parsed from repaired truncated JSON, keys: {list(result.keys())}")
                    return result
                except json.JSONDecodeError as e:
                    logging.warning(f"_extract_json_from_response: Repaired JSON still invalid: {e}")

        # Try 1: Look for ```json code block
        json_match = re.search(r'```json\s*(.*?)\s*```', cleaned_response, re.DOTALL)
        if json_match:
            try:
                result = json.loads(json_match.group(1))
                logging.info(f"_extract_json_from_response: Parsed from ```json block, keys: {list(result.keys())}")
                return result
            except json.JSONDecodeError as e:
                logging.warning(f"_extract_json_from_response: Failed to parse ```json block: {e}")

        # Try 2: Look for ``` code block without json tag
        code_match = re.search(r'```\s*(.*?)\s*```', cleaned_response, re.DOTALL)
        if code_match:
            code_content = code_match.group(1).strip()
            if code_content.startswith('{'):
                try:
                    result = json.loads(code_content)
                    logging.info(f"_extract_json_from_response: Parsed from ``` block, keys: {list(result.keys())}")
                    return result
                except json.JSONDecodeError as e:
                    logging.warning(f"_extract_json_from_response: Failed to parse ``` block: {e}")

        # Try 3: Look for first { to last } (the JSON object)
        first_brace = cleaned_response.find('{')
        last_brace = cleaned_response.rfind('}')
        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            json_str = cleaned_response[first_brace:last_brace + 1]
            try:
                result = json.loads(json_str)
                logging.info(f"_extract_json_from_response: Parsed from braces, keys: {list(result.keys())}")
                return result
            except json.JSONDecodeError as e:
                # Try to fix common JSON issues
                logging.warning(f"_extract_json_from_response: Initial brace parse failed: {e}")

                # Apply multiple fixes in sequence
                fixed_json = json_str

                # Fix 1: Remove trailing commas before ] or }
                fixed_json = re.sub(r',\s*([}\]])', r'\1', fixed_json)

                # Fix 2: Fix missing commas between } and { or ] and [
                fixed_json = re.sub(r'}\s*{', '},{', fixed_json)
                fixed_json = re.sub(r']\s*\[', '],[', fixed_json)

                # Fix 3: Fix missing commas between } and "
                fixed_json = re.sub(r'}\s*"', '},"', fixed_json)
                fixed_json = re.sub(r']\s*"', '],"', fixed_json)

                # Fix 4: Fix newlines inside strings (convert to spaces)
                # This is tricky - need to only fix inside strings
                # For now, just remove literal newlines that aren't escaped
                fixed_json = re.sub(r'(?<!\\)\n', ' ', fixed_json)

                try:
                    result = json.loads(fixed_json)
                    logging.info(f"_extract_json_from_response: Parsed after JSON fixes, keys: {list(result.keys())}")
                    return result
                except json.JSONDecodeError as e2:
                    logging.warning(f"_extract_json_from_response: JSON fixes didn't help: {e2}")

                # Try to find and fix the specific position of the error
                try:
                    # Sometimes the AI includes extra text after the JSON
                    # Try parsing incrementally to find where valid JSON ends
                    for end_pos in range(last_brace, first_brace, -1):
                        if cleaned_response[end_pos] == '}':
                            try:
                                result = json.loads(cleaned_response[first_brace:end_pos + 1])
                                logging.info(f"_extract_json_from_response: Parsed with truncated end, keys: {list(result.keys())}")
                                return result
                            except json.JSONDecodeError:
                                continue
                except Exception:
                    pass

        # Try 4: Direct parse (if entire response is JSON)
        try:
            result = json.loads(cleaned_response)
            logging.info(f"_extract_json_from_response: Parsed directly, keys: {list(result.keys())}")
            return result
        except json.JSONDecodeError as e:
            logging.warning(f"_extract_json_from_response: Direct parse failed: {e}")

        # Failed to parse - log detailed error info
        logging.error(f"_extract_json_from_response: ALL PARSE ATTEMPTS FAILED")
        logging.error(f"Response length: {len(response)}")
        logging.error(f"Response starts with: {response[:200]}")
        logging.error(f"Response ends with: {response[-200:] if len(response) > 200 else response}")

        return {"parse_error": True, "raw_analysis": response, "error_type": "json_parse_failed"}

    def _repair_truncated_json(self, json_str: str) -> Optional[str]:
        """Attempt to repair truncated JSON by closing unclosed brackets/braces."""
        try:
            # Track the nesting level
            stack = []
            in_string = False
            escape_next = False
            last_complete_pos = 0

            for i, char in enumerate(json_str):
                if escape_next:
                    escape_next = False
                    continue

                if char == '\\' and in_string:
                    escape_next = True
                    continue

                if char == '"' and not escape_next:
                    in_string = not in_string
                    continue

                if in_string:
                    continue

                if char == '{':
                    stack.append('}')
                elif char == '[':
                    stack.append(']')
                elif char in '}]':
                    if stack and stack[-1] == char:
                        stack.pop()
                        last_complete_pos = i + 1

                # Track positions after complete key-value pairs
                if char == ',' and not in_string:
                    last_complete_pos = i + 1

            # If we're still in a string, try to close it
            if in_string:
                json_str = json_str + '"'

            # Find the last complete structure and close everything
            # Try to truncate at the last comma and close
            if stack:
                # Find last comma outside of string
                truncate_pos = json_str.rfind(',')
                if truncate_pos > 0:
                    # Truncate at last comma and close all open brackets
                    repaired = json_str[:truncate_pos]
                    repaired += ''.join(reversed(stack))
                    return repaired
                else:
                    # Just close all open brackets
                    return json_str + ''.join(reversed(stack))

            return json_str

        except Exception as e:
            logging.warning(f"_repair_truncated_json failed: {e}")
            return None

    def _call_openai(self, system_prompt: str, user_prompt: str, max_tokens: int = 4000) -> str:
        """Call Azure OpenAI with the given prompts. Handles model-specific parameter differences."""
        if not self.openai_client:
            return "[ERROR: OpenAI client not initialized]"

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]

        # Try different parameter combinations for model compatibility
        # gpt-5.x models may not support temperature or max_tokens
        param_combinations = [
            # Try minimal params first (most compatible with newer models)
            {"model": self.deployment_name, "messages": messages, "max_completion_tokens": max_tokens},
            # Try with max_tokens instead
            {"model": self.deployment_name, "messages": messages, "max_tokens": max_tokens},
            # Try with temperature for older models
            {"model": self.deployment_name, "messages": messages, "max_tokens": max_tokens, "temperature": 0.3},
        ]

        last_error = None
        for params in param_combinations:
            try:
                response = self.openai_client.chat.completions.create(**params)
                return response.choices[0].message.content
            except Exception as e:
                error_msg = str(e).lower()
                # If it's a parameter compatibility error, try next combination
                if "unsupported" in error_msg or "not supported" in error_msg:
                    logging.info(f"Parameter compatibility issue, trying next combination: {e}")
                    last_error = e
                    continue
                # For other errors, fail immediately
                logging.error(f"OpenAI call failed: {e}")
                return f"[ERROR: OpenAI analysis failed: {e}]"

        # All combinations failed
        logging.error(f"All parameter combinations failed. Last error: {last_error}")
        return f"[ERROR: OpenAI analysis failed after trying multiple parameter combinations: {last_error}]"

    def _chunk_text(self, text: str, max_chars: int = 30000) -> List[str]:
        """Split text into chunks for processing large documents."""
        if len(text) <= max_chars:
            return [text]

        chunks = []
        current_chunk = ""

        # Split by paragraphs to maintain context
        paragraphs = text.split('\n\n')

        for para in paragraphs:
            if len(current_chunk) + len(para) < max_chars:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"

        if current_chunk:
            chunks.append(current_chunk.strip())

        return chunks

    def _analyze_full_contract(self, text: str, file_name: str) -> Dict:
        """Perform comprehensive contract analysis from LABEL perspective."""
        system_prompt = """You are an expert contract analyst working for a MAJOR RECORD LABEL (Label perspective).
Your job is to analyze contracts and assess them from the LABEL'S business interests.

This is a BUSINESS ANALYSIS tool for internal use. You are analyzing an ACTUAL contract document.
The document text includes [PAGE N] markers showing where each page begins - use these for references.

Analyze the provided contract and extract structured information.
IMPORTANT: All assessments should be from the LABEL's perspective - what benefits the label, what risks the label faces.

REFERENCE FORMAT: Use the [PAGE N] markers in the text to identify page numbers.
- If you see content after "[PAGE 3]", reference it as "Page 3"
- If the document has visible section/article numbers, include them: "Page 3, Section 4.1"
- If no section numbers are visible, just use the page: "Page 3"
- References help readers locate content but don't need to be formal citations

Return your analysis as a valid JSON object with the following structure:
{
    "contract_type": "type of contract (e.g., Recording Agreement, Licensing Deal, Service Agreement)",
    "parties": [{"name": "party name", "role": "role in contract (Label/Artist/Licensor/etc)", "ref": "Page X"}],
    "effective_date": {"value": "date or null if not found", "ref": "Page X"},
    "term_duration": {"value": "duration description", "ref": "Page X, Section Y"},
    "financial_terms": {
        "advances": {"value": "amount or null", "ref": "Page X, Section Y"},
        "royalty_rates": {"value": "rates description", "ref": "Page X, Section Y"},
        "payment_schedule": {"value": "description or null", "ref": "Page X"},
        "label_investment": {"value": "total label financial commitment", "ref": "Page X"},
        "recoupment_terms": {"value": "how label recoups investment", "ref": "Page X, Section Y"},
        "other_payments": [{"description": "...", "ref": "Page X"}]
    },
    "rights_secured": [{"right": "description", "scope": "scope/territory", "exclusivity": true/false, "duration": "how long label holds rights", "ref": "Page X, Section Y"}],
    "label_protections": [{"clause": "protection description", "ref": "Page X, Section Y"}],
    "artist_obligations": [{"obligation": "what artist must do", "deadline": "when", "consequence_if_breached": "label remedy", "ref": "Page X, Section Y"}],
    "label_obligations": [{"obligation": "what label must do", "financial_impact": "cost to label", "ref": "Page X, Section Y"}],
    "termination_clauses": [{"trigger": "what triggers termination", "who_can_trigger": "label/artist/either", "label_impact": "effect on label", "ref": "Page X, Section Y"}],
    "key_dates": [{"event": "description", "date": "date or relative timing", "ref": "Page X"}],
    "artist_favorable_terms": [{"term": "term that favors artist MORE than industry standard", "concern": "why this is a concern for label", "ref": "Page X, Section Y"}],
    "missing_label_protections": ["standard label protections that appear to be missing"],
    "overall_assessment": "2-3 sentence assessment from LABEL's perspective - is this deal favorable to the label?"
}

Be thorough but concise. Extract actual values from the contract text.
Remember: You work for the LABEL. Artist-favorable terms are potential concerns.
ALWAYS include page/section references for every finding.

CRITICAL INSTRUCTIONS - YOU MUST FOLLOW THESE:
1. Output ONLY the JSON object - no other text before or after
2. Do NOT ask any questions - just analyze and output JSON immediately
3. Do NOT ask for confirmation or clarification - proceed with analysis
4. Do NOT mention length limits or offer to split output - just output the JSON
5. Start your response with { and end with } - nothing else
6. If information is missing, use null or "Not specified" - do not ask about it

Your response must be valid JSON starting with { and ending with }."""

        # Handle large documents by chunking
        chunks = self._chunk_text(text)

        if len(chunks) == 1:
            user_prompt = f"Analyze this contract:\n\n{text}"
            # Use larger max_tokens for full contract analysis to accommodate detailed JSON
            response = self._call_openai(system_prompt, user_prompt, max_tokens=16000)
        else:
            # For large documents, analyze chunks and synthesize
            chunk_analyses = []
            for i, chunk in enumerate(chunks):
                user_prompt = f"Analyze this section (Part {i+1} of {len(chunks)}) of a contract:\n\n{chunk}"
                chunk_response = self._call_openai(system_prompt, user_prompt, max_tokens=2000)
                chunk_analyses.append(chunk_response)

            # Synthesize the chunks
            synthesis_prompt = """You are synthesizing multiple partial analyses of a single contract.
Combine these analyses into one comprehensive JSON structure, removing duplicates and resolving any conflicts.

CRITICAL: Output ONLY valid JSON. Start with { and end with }. No other text.
Do NOT ask questions. Do NOT offer options. Just output the combined JSON immediately."""

            user_prompt = f"Combine these partial contract analyses:\n\n" + "\n\n---\n\n".join(chunk_analyses)
            response = self._call_openai(synthesis_prompt, user_prompt)

        # Parse the JSON response with retry logic for empty results
        essential_keys = ['contract_type', 'parties', 'rights_secured', 'financial_terms']
        max_retries = 3

        for attempt in range(max_retries):
            analysis = self._extract_json_from_response(response)

            # Check if analysis has essential keys (not just empty or error)
            has_essential_data = any(
                key in analysis and analysis[key]
                for key in essential_keys
            )

            if has_essential_data or analysis.get('parse_error'):
                # Either we got good data or a clear error - don't retry
                break

            if attempt < max_retries - 1:
                logging.warning(f"Full analysis returned empty (attempt {attempt + 1}/{max_retries}). Retrying...")
                # Retry with more forceful prompt
                retry_prompt = f"""IMPORTANT: Your previous response was empty or incomplete.
You MUST output a complete JSON object with contract analysis.

Analyze this contract NOW and output ONLY the JSON (start with {{ end with }}):

{text[:40000]}"""
                response = self._call_openai(system_prompt, retry_prompt, max_tokens=16000)

        # DEBUG: Enhanced logging for analysis result
        logging.info("=" * 40)
        logging.info("_analyze_full_contract - RESULT ANALYSIS")
        logging.info(f"Full analysis completed with keys: {list(analysis.keys())}")
        logging.info(f"Raw response length: {len(response) if response else 0}")
        logging.info(f"Raw response preview: {response[:500] if response else 'None'}")

        if analysis.get('parse_error'):
            logging.error(f"_analyze_full_contract - JSON PARSE FAILED!")
            logging.error(f"Raw analysis: {analysis.get('raw_analysis', '')[:1000]}")

        # Log specific key values
        logging.info(f"  - contract_type: {analysis.get('contract_type', 'MISSING')}")
        logging.info(f"  - parties: {len(analysis.get('parties', []))} found")
        logging.info(f"  - financial_terms: {type(analysis.get('financial_terms')).__name__}")
        logging.info(f"  - rights_secured: {len(analysis.get('rights_secured', []))} found")

        if not any(key in analysis for key in essential_keys):
            logging.warning(f"Analysis may be incomplete - no essential keys found!")
        logging.info("=" * 40)

        analysis["_metadata"] = {
            "file_name": file_name,
            "analyzed_at": datetime.now().isoformat(),
            "text_length": len(text),
            "chunks_processed": len(chunks),
            "retry_attempts": attempt + 1
        }
        return analysis

    def _extract_specific_clauses(self, text: str, clause_types: List[str]) -> Dict:
        """Extract specific types of clauses from the contract."""
        clause_descriptions = {
            "financial": "All financial terms including advances, royalties, payments, fees, revenue sharing, expenses",
            "rights": "All rights granted or reserved including intellectual property, licensing, usage rights, exclusivity",
            "obligations": "All obligations and duties of each party, deliverables, performance requirements",
            "termination": "Termination conditions, notice periods, breach definitions, consequences of termination",
            "exclusivity": "Exclusivity clauses, non-compete provisions, first refusal rights",
            "territory": "Geographic scope, territory definitions, regional limitations",
            "duration": "Term length, renewal options, extension conditions, effective dates"
        }

        types_to_extract = [ct for ct in clause_types if ct in clause_descriptions]
        if not types_to_extract:
            types_to_extract = list(clause_descriptions.keys())

        extraction_details = "\n".join([f"- {ct}: {clause_descriptions[ct]}" for ct in types_to_extract])

        system_prompt = f"""You are a contract clause extraction specialist.
Extract the following types of clauses from the contract:

{extraction_details}

CRITICAL: For EVERY clause extracted, you MUST include a document reference:
- "ref": "Page X, Section Y" or "ref": "Page X, Article Y"
The document text includes [PAGE N] markers - use these for exact page numbers.

Return a JSON object where each key is a clause type and the value is an array of extracted clauses.
Each clause should have: "text" (the clause text or summary - keep under 200 words), "ref" (Page X, Section Y), "key_points" (2-3 bullet points).

Example format:
{{
    "financial": [
        {{"text": "...", "ref": "Page 3, Section 4.1", "key_points": ["Advance of $X", "Royalty rate of Y%"]}}
    ]
}}

ALWAYS include page and section references for every extracted clause.

CRITICAL INSTRUCTIONS - YOU MUST FOLLOW THESE:
1. Output ONLY the JSON object - no other text before or after
2. Do NOT ask any questions - just extract and output JSON immediately
3. Do NOT ask for confirmation or clarification - proceed with extraction
4. Do NOT mention length limits or offer to split output - output the JSON directly
5. Start your response with {{ and end with }} - nothing else
6. If a clause type has no matches, use an empty array: "type": []
7. Keep each clause summary under 200 words - summarize if needed
8. Extract UP TO 3-5 most important clauses per category (prioritize key terms)

Your response must be valid JSON starting with {{ and ending with }}."""

        # For large documents, process only the most relevant text
        text_to_process = text[:35000]  # Reduced limit for better processing

        user_prompt = f"Extract clauses from this contract. Output ONLY valid JSON:\n\n{text_to_process}"
        response = self._call_openai(system_prompt, user_prompt, max_tokens=6000)

        # Add logging for debugging
        logging.info(f"_extract_specific_clauses - Response length: {len(response) if response else 0}")

        result = self._extract_json_from_response(response)

        # If parse failed, try with retry
        if result.get('parse_error'):
            logging.warning("_extract_specific_clauses - Initial parse failed, retrying...")
            retry_prompt = f"""IMPORTANT: Output ONLY valid JSON. No other text. Start with {{ end with }}.

Extract clauses from this contract into JSON format:
{text_to_process[:25000]}"""
            response = self._call_openai(system_prompt, retry_prompt, max_tokens=6000)
            result = self._extract_json_from_response(response)

        return result

    def _generate_summary(self, text: str, summary_type: str, audience: str) -> Dict:
        """Generate a summary tailored to the audience - FROM LABEL PERSPECTIVE."""
        audience_instructions = {
            "legal": "Use precise legal terminology. Include specific clause references. Highlight legal risks to the LABEL and compliance considerations.",
            "business": "Focus on commercial terms and business implications FOR THE LABEL. Emphasize label ROI, recoupment timeline, and operational impacts to the label.",
            "executive": "Provide a high-level overview for LABEL C-suite. Focus on label investment, rights secured, strategic value, and key risks to the label. Keep it concise."
        }

        length_instructions = {
            "executive": "Provide a brief 150-200 word summary with key bullet points.",
            "detailed": "Provide a comprehensive 400-500 word summary covering all major aspects.",
            "legal": "Provide a thorough legal summary of 300-400 words with specific clause references."
        }

        system_prompt = f"""You are a contract summarization expert working for a MAJOR RECORD LABEL.
Your summaries are FOR LABEL EXECUTIVES and should reflect the LABEL's interests and perspective.

{audience_instructions.get(audience, audience_instructions['business'])}
{length_instructions.get(summary_type, length_instructions['detailed'])}

IMPORTANT: Frame everything from the label's business perspective:
- "Label investment" instead of "artist advance"
- "Rights secured by label" instead of "rights granted"
- "Artist-favorable terms" = potential concerns for the label
- Risk assessment = risks TO THE LABEL

CRITICAL: Include document references (Page X, Section Y) for key claims.
The document text includes [PAGE N] markers - use these for exact page numbers.

CRITICAL: The recommendation MUST align with the risk level:
- LOW risk: "Proceed" or "Proceed as drafted"
- MEDIUM risk: "Proceed with caution; consider negotiating [specific terms]"
- HIGH risk: "Do not proceed without changes to [specific critical terms]" or "Renegotiate [specific issues] before proceeding"

Return a JSON object with:
{{
    "summary": "the main summary text FROM LABEL PERSPECTIVE",
    "key_points": [{{"point": "bullet point text", "ref": "Page X, Section Y"}}],
    "label_investment_total": {{"value": "total financial commitment from label", "ref": "Page X"}},
    "rights_secured": {{"value": "summary of rights label obtains", "ref": "Page X, Section Y"}},
    "critical_dates": [{{"event": "...", "date": "...", "ref": "Page X"}}],
    "action_items": ["any actions needed by label team"],
    "risk_level": "low/medium/high (risk TO THE LABEL)",
    "artist_leverage_concerns": [{{"concern": "term giving artist unusual leverage", "ref": "Page X, Section Y"}}],
    "recommendation": "recommendation aligned with risk level - if HIGH risk, must specify required changes before proceeding"
}}

ALWAYS include page/section references for key findings.

IMPORTANT: Do NOT ask clarifying questions. Do NOT ask for confirmation. Just execute the summary and return the JSON."""

        user_prompt = f"Summarize this contract for a {audience} audience:\n\n{text[:40000]}"
        response = self._call_openai(system_prompt, user_prompt)
        return self._extract_json_from_response(response)

    def _identify_risks(self, text: str) -> Dict:
        """Identify risks TO THE LABEL and deviations from standard terms."""
        system_prompt = """You are a contract risk analyst working for a MAJOR RECORD LABEL.
Your job is to identify risks TO THE LABEL, not to the artist.

IMPORTANT PERSPECTIVE:
- Artist-favorable terms = RISKS to the label (higher costs, less control, early reversion)
- High advances/royalties = FINANCIAL RISK to label
- Strong artist termination rights = OPERATIONAL RISK to label
- Early master reversion = ASSET RISK to label
- Creative control for artist = COMMERCIAL RISK to label
- Non-recoupable payments = DIRECT COST to label

CRITICAL: For EVERY risk identified, you MUST include a document reference:
- "ref": "Page X, Section Y" or "ref": "Page X, Article Y"
The document text includes [PAGE N] markers - use these for exact page numbers.

Analyze the contract for RISKS TO THE LABEL:
1. Financial exposure (high advances, guaranteed payments, non-recoupable costs)
2. Asset risks (early master reversion, limited rights duration, territory restrictions)
3. Operational risks (artist approval requirements, key man clauses, delivery delays)
4. Revenue risks (high royalty rates, favorable streaming splits, limited 360 participation)
5. Legal/compliance risks (regulatory, indemnification gaps)
6. Competitive risks (artist leverage, termination options)

Return a JSON object:
{
    "overall_risk_level": "low/medium/high (RISK TO LABEL)",
    "risk_score": 1-100 (higher = worse for label),
    "label_financial_exposure": "total potential label investment/loss",
    "risks": [
        {
            "category": "financial/asset/operational/revenue/legal/competitive",
            "severity": "low/medium/high/critical",
            "description": "what the risk TO THE LABEL is",
            "ref": "Page X, Section Y",
            "label_impact": "specific impact on label operations/finances",
            "recommendation": "how label should address this in negotiation"
        }
    ],
    "artist_favorable_terms": [
        {"term": "what favors the artist", "industry_standard": "what's typical", "label_impact": "why this hurts the label", "ref": "Page X, Section Y"}
    ],
    "missing_label_protections": ["standard label protections not found in this contract"],
    "negotiation_priorities": [{"priority": "item label should push back on", "ref": "Page X, Section Y"}],
    "deal_breakers": [{"issue": "term that may be unacceptable", "ref": "Page X, Section Y"}],
    "summary": "2-3 sentence risk summary FROM LABEL'S PERSPECTIVE"
}

ALWAYS cite the specific page and section for every risk and concern.

IMPORTANT: Do NOT ask clarifying questions. Do NOT ask for confirmation. Just execute the analysis and return the JSON."""

        user_prompt = f"Analyze risks in this contract:\n\n{text[:50000]}"
        response = self._call_openai(system_prompt, user_prompt)

        # DEBUG: Log the raw response before parsing
        logging.info(f"_identify_risks - Raw response length: {len(response) if response else 0}")
        logging.info(f"_identify_risks - Raw response preview: {response[:1000] if response else 'EMPTY'}")

        result = self._extract_json_from_response(response)

        # DEBUG: Log the parsed result
        logging.info(f"_identify_risks - Parsed result keys: {list(result.keys()) if isinstance(result, dict) else 'NOT DICT'}")
        if isinstance(result, dict) and result.get('parse_error'):
            logging.error(f"_identify_risks - JSON PARSE FAILED!")

        return result

    def _compare_contracts(self, text_a: str, text_b: str, name_a: str, name_b: str) -> Dict:
        """Compare two contracts using sectioned analysis for better coverage."""

        # For any substantial contracts, use sectioned comparison to avoid truncation
        # Sectioned mode analyzes each area (financial, rights, etc.) separately
        # This provides better coverage and avoids the model refusing due to incomplete text
        total_length = len(text_a) + len(text_b)
        max_single_contract = max(len(text_a), len(text_b))

        # Use sectioned mode if total > 30k OR if either contract alone is > 20k
        # This ensures we don't truncate important contract content
        use_sectioned = total_length > 30000 or max_single_contract > 20000

        if use_sectioned:
            return self._compare_contracts_sectioned(text_a, text_b, name_a, name_b)

        system_prompt = """You are a contract analysis assistant performing an EDUCATIONAL comparison of two recording agreements.
This is a BUSINESS ANALYSIS exercise for training purposes, NOT legal advice.

Your task: Compare these two contracts and identify factual differences in their terms.
Analyze which contract has more favorable terms from a BUSINESS perspective (lower financial commitments, stronger protections, better rights retention).

When comparing terms, note which contract (A or B) has:
- Lower advance/payment obligations
- Longer rights retention periods
- Broader territorial coverage
- More comprehensive protections
- Clearer deliverable requirements

CRITICAL: Include document references for BOTH contracts in the format:
- "ref_a": "Page X, Section Y" (for Contract A)
- "ref_b": "Page X, Section Y" (for Contract B)
The document text includes [PAGE N] markers - use these for exact page numbers.

Compare these two contracts and identify:
1. Key differences in financial terms, rights, obligations
2. Which contract has more protective clauses (from a business standpoint)
3. Which contract has higher financial exposure
4. Notable terms that differ significantly between the two

Return a JSON object:
{
    "similarity_score": 0-100,
    "contract_types_match": true/false,
    "more_label_favorable": "a/b/neutral (which contract has more protective business terms)",
    "key_differences": [
        {
            "aspect": "what's being compared",
            "contract_a": "terms in first contract",
            "contract_b": "terms in second contract",
            "ref_a": "Page X, Section Y",
            "ref_b": "Page X, Section Y",
            "label_preference": "a/b (which has stronger business protections)",
            "label_impact": "business significance of this difference"
        }
    ],
    "unique_to_a": [{"clause": "clause description", "ref": "Page X, Section Y"}],
    "unique_to_b": [{"clause": "clause description", "ref": "Page X, Section Y"}],
    "financial_comparison": {
        "contract_a": {"label_investment": "total financial commitment", "royalty_exposure": "royalty rates", "ref": "Page X"},
        "contract_b": {"label_investment": "total financial commitment", "royalty_exposure": "royalty rates", "ref": "Page X"},
        "lower_label_cost": "a/b",
        "better_label_margin": "a/b"
    },
    "rights_comparison": {
        "contract_a": {"rights_duration": "...", "territory": "...", "reversion": "...", "ref": "Page X, Section Y"},
        "contract_b": {"rights_duration": "...", "territory": "...", "reversion": "...", "ref": "Page X, Section Y"},
        "stronger_label_rights": "a/b"
    },
    "risk_comparison": {
        "contract_a_risk_level": "low/medium/high (financial/operational risk level)",
        "contract_b_risk_level": "low/medium/high (financial/operational risk level)",
        "lower_label_risk": "a/b"
    },
    "overall_assessment": "factual summary of which contract has more favorable business terms and key differences",
    "recommended_standard_terms": ["notable terms from either contract worth considering"]
}

ALWAYS include page/section references from both contracts for every comparison point.

IMPORTANT: This is an educational analysis. Provide factual comparisons. Do NOT ask clarifying questions. Just execute the comparison and return the JSON."""

        user_prompt = f"""Compare these two contracts:

=== CONTRACT A: {name_a} ===
{text_a[:30000]}

=== CONTRACT B: {name_b} ===
{text_b[:30000]}"""

        response = self._call_openai(system_prompt, user_prompt, max_tokens=6000)
        result = self._extract_json_from_response(response)
        result["_metadata"] = {
            "contract_a": name_a,
            "contract_b": name_b,
            "compared_at": datetime.now().isoformat(),
            "comparison_mode": "standard"
        }
        return result

    def _compare_contracts_sectioned(self, text_a: str, text_b: str, name_a: str, name_b: str) -> Dict:
        """Compare large contracts by analyzing sections separately then synthesizing."""
        logging.info(f"Using sectioned comparison for large contracts: {name_a} vs {name_b}")

        # Define comparison sections
        sections = {
            "financial": {
                "focus": "advances, royalties, recoupment, payment schedules, 360 terms, merchandise, touring splits",
                "analysis_criteria": "Compare total financial commitments, royalty rates, recoupment structures"
            },
            "rights": {
                "focus": "master ownership, duration of rights, territory, exclusivity, reversion triggers, publishing",
                "analysis_criteria": "Compare scope and duration of rights granted, territorial coverage, reversion terms"
            },
            "obligations": {
                "focus": "delivery requirements, album commitments, promotional obligations, key man clauses",
                "analysis_criteria": "Compare deliverable requirements, commitment levels, operational obligations"
            },
            "termination_risk": {
                "focus": "termination triggers, exit clauses, breach remedies, force majeure, key man provisions",
                "analysis_criteria": "Compare termination conditions, exit mechanisms, breach remedies"
            }
        }

        section_results = {}

        for section_name, section_info in sections.items():
            section_prompt = f"""You are a contract analysis assistant performing an EDUCATIONAL comparison of two recording agreements.
Focus ONLY on {section_name.upper()} terms in these two contracts.

FOCUS AREAS: {section_info['focus']}
ANALYSIS CRITERIA: {section_info['analysis_criteria']}

Extract and compare {section_name} terms from both contracts factually.

Return a JSON object with this EXACT structure:
{{
    "section": "{section_name}",
    "contract_a_terms": {{
        "summary": "brief summary of {section_name} terms in Contract A",
        "key_values": ["specific values/terms found"],
        "refs": ["Page X, Section Y"]
    }},
    "contract_b_terms": {{
        "summary": "brief summary of {section_name} terms in Contract B",
        "key_values": ["specific values/terms found"],
        "refs": ["Page X, Section Y"]
    }},
    "differences": [
        {{
            "aspect": "specific term being compared",
            "contract_a": "value/term in A",
            "contract_b": "value/term in B",
            "label_preference": "a/b/neutral (which has more favorable business terms)",
            "reason": "factual explanation of the difference"
        }}
    ],
    "section_winner": "a/b/neutral (which contract has more favorable terms in this section)",
    "section_assessment": "1-2 sentence factual assessment of key differences"
}}

IMPORTANT: This is educational analysis. Do NOT ask questions. Just analyze and return the JSON."""

            user_prompt = f"""Compare {section_name.upper()} terms:

=== CONTRACT A: {name_a} ===
{text_a[:50000]}

=== CONTRACT B: {name_b} ===
{text_b[:50000]}"""

            response = self._call_openai(section_prompt, user_prompt, max_tokens=3000)
            section_results[section_name] = self._extract_json_from_response(response)

        # Synthesize all section results into final comparison
        synthesis_prompt = """You are an educational contract analysis assistant synthesizing sectioned comparison results.

Based on the section-by-section analysis provided, create a comprehensive factual comparison summary.

Return a JSON object:
{
    "similarity_score": 0-100,
    "contract_types_match": true/false,
    "more_label_favorable": "a/b/neutral (which has more favorable business terms overall)",
    "section_winners": {
        "financial": "a/b/neutral",
        "rights": "a/b/neutral",
        "obligations": "a/b/neutral",
        "termination_risk": "a/b/neutral"
    },
    "key_differences": [
        {
            "aspect": "term being compared",
            "contract_a": "value in A",
            "contract_b": "value in B",
            "label_preference": "a/b (which has more favorable terms)",
            "label_impact": "business significance of this difference"
        }
    ],
    "financial_comparison": {
        "contract_a": {"label_investment": "total commitment amount", "royalty_exposure": "royalty rates"},
        "contract_b": {"label_investment": "total commitment amount", "royalty_exposure": "royalty rates"},
        "lower_label_cost": "a/b",
        "better_label_margin": "a/b"
    },
    "rights_comparison": {
        "contract_a": {"rights_duration": "...", "territory": "...", "reversion": "..."},
        "contract_b": {"rights_duration": "...", "territory": "...", "reversion": "..."},
        "stronger_label_rights": "a/b"
    },
    "risk_comparison": {
        "contract_a_risk_level": "low/medium/high",
        "contract_b_risk_level": "low/medium/high",
        "lower_label_risk": "a/b"
    },
    "deal_breakers": [{"contract": "a/b", "issue": "notable concern"}],
    "overall_assessment": "2-3 sentence factual summary of key differences between the contracts",
    "recommended_standard_terms": ["notable terms from either contract"]
}

IMPORTANT: This is educational analysis. Provide factual comparisons only."""

        synthesis_user = f"""Synthesize these section comparisons:

{json.dumps(section_results, indent=2, default=str)}

Contract A: {name_a}
Contract B: {name_b}"""

        synthesis_response = self._call_openai(synthesis_prompt, synthesis_user, max_tokens=4000)
        result = self._extract_json_from_response(synthesis_response)

        # Add section details and metadata
        result["_section_details"] = section_results
        result["_metadata"] = {
            "contract_a": name_a,
            "contract_b": name_b,
            "compared_at": datetime.now().isoformat(),
            "comparison_mode": "sectioned",
            "sections_analyzed": list(sections.keys())
        }

        return result

    def perform(self, **kwargs) -> str:
        """Execute contract analysis action."""
        try:
            action = kwargs.get('action', 'list_contracts')
            contract_name = kwargs.get('contract_name')
            contract_name_b = kwargs.get('contract_name_b')
            clause_types = kwargs.get('clause_types', [])
            summary_type = kwargs.get('summary_type', 'detailed')
            audience = kwargs.get('audience', 'business')

            # List contracts
            if action == 'list_contracts':
                files = self._list_files_in_folder()
                return json.dumps({
                    "status": "success",
                    "action": "list_contracts",
                    "contracts_folder": self.contracts_folder,
                    "files": files,
                    "count": len(files),
                    "supported_formats": ["pdf", "docx", "txt"],
                    "usage": "Use contract_name parameter with the file name to analyze"
                }, indent=2)

            # All other actions require a contract name
            if not contract_name:
                return json.dumps({
                    "status": "error",
                    "message": "contract_name is required for this action",
                    "available_contracts": self._list_files_in_folder()
                }, indent=2)

            # Read the contract
            file_path = f"{self.contracts_folder}/{contract_name}"
            content = self._read_file_content(file_path)

            if not content:
                return json.dumps({
                    "status": "error",
                    "message": f"Could not read contract: {contract_name}",
                    "path_tried": file_path,
                    "available_contracts": self._list_files_in_folder()
                }, indent=2)

            # Extract text
            text = self._extract_text(file_path, content)
            if text.startswith("[ERROR"):
                return json.dumps({
                    "status": "error",
                    "message": text
                }, indent=2)

            # Execute the requested action
            if action == 'full_workup':
                # Comprehensive analysis: runs everything in one go
                logging.info(f"Running full workup on {contract_name}")

                # 1. Full contract analysis
                analysis = self._analyze_full_contract(text, contract_name)

                # 2. Risk identification (run first - this is the authoritative risk source)
                risks = self._identify_risks(text)

                # 3. Executive summary for business audience
                summary = self._generate_summary(text, 'executive', 'business')

                # 4. Extract key clauses (all types)
                all_clause_types = ['financial', 'rights', 'obligations', 'termination', 'exclusivity', 'territory', 'duration']
                clauses = self._extract_specific_clauses(text, all_clause_types)

                # Synchronize risk levels - use risk assessment as authoritative source
                # This ensures consistency throughout the PDF report
                authoritative_risk_level = risks.get('overall_risk_level', 'unknown')
                authoritative_risk_score = risks.get('risk_score', 'N/A')
                if isinstance(summary, dict):
                    summary['risk_level'] = authoritative_risk_level
                    summary['risk_score'] = authoritative_risk_score

                # Compile full report (include contract text for PDF with clickable references)
                full_report = {
                    "contract": contract_name,
                    "analyzed_at": datetime.now().isoformat(),
                    "executive_summary": summary,
                    "full_analysis": analysis,
                    "risk_assessment": risks,
                    "extracted_clauses": clauses,
                    "text_length": len(text),
                    "_contract_text": text  # Include for PDF generation with clickable refs
                }

                # Save the analysis report to Azure storage
                save_result = self._save_analysis_report(contract_name, full_report)

                # Build concise chat response (fits on one screen)
                risk_level = risks.get('overall_risk_level', 'Unknown').upper()
                risk_score = risks.get('risk_score', 'N/A')
                risk_summary = risks.get('summary', '')

                # Get top 3 risks with references
                top_risks = []
                for r in risks.get('risks', [])[:3]:
                    if isinstance(r, dict):
                        ref = r.get('ref', 'N/A')
                        desc = r.get('description', '')[:80]
                        severity = r.get('severity', '').upper()
                        top_risks.append(f"- [{severity}] {desc} (Ref: {ref})")

                # Get key financial terms with references
                fin_terms = analysis.get('financial_terms', {})
                advances = fin_terms.get('advances', {})
                adv_val = advances.get('value', 'N/A') if isinstance(advances, dict) else advances
                adv_ref = advances.get('ref', '') if isinstance(advances, dict) else ''

                royalties = fin_terms.get('royalty_rates', {})
                roy_val = royalties.get('value', 'N/A') if isinstance(royalties, dict) else royalties
                roy_ref = royalties.get('ref', '') if isinstance(royalties, dict) else ''

                # Build the chat summary (short, fits one screen)
                chat_summary = {
                    "headline": f"Analysis Complete: {contract_name}",
                    "risk_level": risk_level,
                    "risk_score": f"{risk_score}/100",
                    "key_findings": [
                        f"Advance: {adv_val}" + (f" (Ref: {adv_ref})" if adv_ref else ""),
                        f"Royalty: {roy_val}" + (f" (Ref: {roy_ref})" if roy_ref else ""),
                    ],
                    "top_risks": top_risks,
                    "recommendation": summary.get('recommendation', '') if isinstance(summary, dict) else '',
                    "full_report": {
                        "message": "Full analysis with all details saved to PDF report:",
                        "download_url": save_result.get('download_url', 'Report generation failed'),
                        "report_name": save_result.get('report_name', ''),
                        "size_kb": save_result.get('size_kb', 0)
                    }
                }

                return json.dumps({
                    "status": "success",
                    "action": "full_workup",
                    "contract": contract_name,
                    "chat_response": chat_summary,
                    "_full_data": {
                        "executive_summary": summary,
                        "full_analysis": analysis,
                        "risk_assessment": risks,
                        "extracted_clauses": clauses,
                        "report_saved": save_result
                    },
                    "_metadata": {
                        "analyzed_at": datetime.now().isoformat(),
                        "file_name": contract_name,
                        "text_length": len(text)
                    }
                }, indent=2)

            elif action == 'analyze_contract':
                result = self._analyze_full_contract(text, contract_name)
                return json.dumps({
                    "status": "success",
                    "action": "analyze_contract",
                    "contract": contract_name,
                    "analysis": result
                }, indent=2)

            elif action == 'extract_clauses':
                result = self._extract_specific_clauses(text, clause_types)

                # Build concise chat summary
                clause_summary = []
                for clause_type, clauses in result.items():
                    if isinstance(clauses, list) and len(clauses) > 0:
                        # Get first clause of each type with its reference
                        first_clause = clauses[0]
                        if isinstance(first_clause, dict):
                            clause_summary.append({
                                "type": clause_type,
                                "count": len(clauses),
                                "sample": first_clause.get('text', '')[:100] + "..." if len(first_clause.get('text', '')) > 100 else first_clause.get('text', ''),
                                "ref": first_clause.get('ref', 'N/A')
                            })

                chat_summary = {
                    "headline": f"Clauses Extracted: {contract_name}",
                    "clause_types_found": list(result.keys()) if isinstance(result, dict) else [],
                    "summary": clause_summary
                }

                return json.dumps({
                    "status": "success",
                    "action": "extract_clauses",
                    "contract": contract_name,
                    "clause_types_requested": clause_types or "all",
                    "chat_response": chat_summary,
                    "_full_data": {"extractions": result}
                }, indent=2)

            elif action == 'summarize_contract':
                result = self._generate_summary(text, summary_type, audience)

                # Build concise chat response
                summary_text = result.get('summary', '') if isinstance(result, dict) else str(result)

                # Extract key points with refs
                key_points = []
                for pt in (result.get('key_points', []) if isinstance(result, dict) else [])[:4]:
                    if isinstance(pt, dict):
                        key_points.append({
                            "point": pt.get('point', ''),
                            "ref": pt.get('ref', 'N/A')
                        })
                    else:
                        key_points.append({"point": str(pt), "ref": "N/A"})

                chat_summary = {
                    "headline": f"Summary: {contract_name}",
                    "risk_level": result.get('risk_level', 'N/A') if isinstance(result, dict) else 'N/A',
                    "summary": summary_text[:300] + "..." if len(summary_text) > 300 else summary_text,
                    "key_points": key_points,
                    "recommendation": result.get('recommendation', '') if isinstance(result, dict) else ''
                }

                return json.dumps({
                    "status": "success",
                    "action": "summarize_contract",
                    "contract": contract_name,
                    "summary_type": summary_type,
                    "audience": audience,
                    "chat_response": chat_summary,
                    "_full_data": {"result": result}
                }, indent=2)

            elif action == 'identify_risks':
                result = self._identify_risks(text)

                # Build concise chat response
                risk_level = result.get('overall_risk_level', 'Unknown').upper()
                risk_score = result.get('risk_score', 'N/A')

                # Get top 3 risks with references
                top_risks = []
                for r in result.get('risks', [])[:3]:
                    if isinstance(r, dict):
                        ref = r.get('ref', 'N/A')
                        desc = r.get('description', '')[:100]
                        severity = r.get('severity', '').upper()
                        top_risks.append({
                            "severity": severity,
                            "description": desc,
                            "ref": ref
                        })

                # Get deal breakers with references
                deal_breakers = []
                for db in result.get('deal_breakers', [])[:2]:
                    if isinstance(db, dict):
                        deal_breakers.append({
                            "issue": db.get('issue', ''),
                            "ref": db.get('ref', 'N/A')
                        })

                chat_summary = {
                    "headline": f"Risk Analysis: {contract_name}",
                    "risk_level": risk_level,
                    "risk_score": f"{risk_score}/100",
                    "financial_exposure": result.get('label_financial_exposure', 'N/A'),
                    "top_risks": top_risks,
                    "deal_breakers": deal_breakers,
                    "summary": result.get('summary', '')
                }

                return json.dumps({
                    "status": "success",
                    "action": "identify_risks",
                    "contract": contract_name,
                    "chat_response": chat_summary,
                    "_full_data": {"risk_analysis": result}
                }, indent=2)

            elif action == 'compare_contracts':
                if not contract_name_b:
                    return json.dumps({
                        "status": "error",
                        "message": "contract_name_b is required for comparison",
                        "available_contracts": self._list_files_in_folder()
                    }, indent=2)

                # Read second contract
                file_path_b = f"{self.contracts_folder}/{contract_name_b}"
                content_b = self._read_file_content(file_path_b)

                if not content_b:
                    return json.dumps({
                        "status": "error",
                        "message": f"Could not read second contract: {contract_name_b}"
                    }, indent=2)

                text_b = self._extract_text(file_path_b, content_b)
                result = self._compare_contracts(text, text_b, contract_name, contract_name_b)

                # Build concise chat response
                more_favorable = result.get('more_label_favorable', 'neutral')
                winner = contract_name if more_favorable == 'a' else (contract_name_b if more_favorable == 'b' else 'Neither')

                # Get comparison mode from metadata
                metadata = result.get('_metadata', {})
                comparison_mode = metadata.get('comparison_mode', 'standard')

                # Get top key differences with refs
                key_diffs = []
                for diff in result.get('key_differences', [])[:5]:
                    if isinstance(diff, dict):
                        key_diffs.append({
                            "aspect": diff.get('aspect', ''),
                            "contract_a": diff.get('contract_a', '')[:60],
                            "contract_b": diff.get('contract_b', '')[:60],
                            "preference": diff.get('label_preference', '').upper(),
                            "impact": diff.get('label_impact', '')[:100] if diff.get('label_impact') else ''
                        })

                # Financial comparison
                fin_comp = result.get('financial_comparison', {})
                financial_summary = {
                    "contract_a_investment": fin_comp.get('contract_a', {}).get('label_investment', 'N/A') if isinstance(fin_comp.get('contract_a'), dict) else 'N/A',
                    "contract_b_investment": fin_comp.get('contract_b', {}).get('label_investment', 'N/A') if isinstance(fin_comp.get('contract_b'), dict) else 'N/A',
                    "lower_cost": fin_comp.get('lower_label_cost', 'N/A'),
                    "better_margin": fin_comp.get('better_label_margin', 'N/A')
                }

                # Rights comparison
                rights_comp = result.get('rights_comparison', {})
                rights_summary = {
                    "contract_a_duration": rights_comp.get('contract_a', {}).get('rights_duration', 'N/A') if isinstance(rights_comp.get('contract_a'), dict) else 'N/A',
                    "contract_b_duration": rights_comp.get('contract_b', {}).get('rights_duration', 'N/A') if isinstance(rights_comp.get('contract_b'), dict) else 'N/A',
                    "stronger_rights": rights_comp.get('stronger_label_rights', 'N/A')
                }

                # Risk comparison
                risk_comp = result.get('risk_comparison', {})

                # Section winners (for sectioned mode)
                section_winners = result.get('section_winners', {})

                # Deal breakers
                deal_breakers = []
                for db in result.get('deal_breakers', [])[:3]:
                    if isinstance(db, dict):
                        deal_breakers.append({
                            "contract": db.get('contract', ''),
                            "issue": db.get('issue', '')[:100]
                        })

                chat_summary = {
                    "headline": f"CONTRACT COMPARISON: {contract_name} vs {contract_name_b}",
                    "comparison_mode": comparison_mode,
                    "overall_winner": {
                        "more_favorable_to_label": winner,
                        "verdict": "Contract A" if more_favorable == 'a' else ("Contract B" if more_favorable == 'b' else "Neutral - Neither clearly better")
                    },
                    "section_breakdown": {
                        "financial": section_winners.get('financial', fin_comp.get('lower_label_cost', 'N/A')),
                        "rights": section_winners.get('rights', rights_comp.get('stronger_label_rights', 'N/A')),
                        "obligations": section_winners.get('obligations', 'N/A'),
                        "termination_risk": section_winners.get('termination_risk', risk_comp.get('lower_label_risk', 'N/A'))
                    },
                    "financial_comparison": financial_summary,
                    "rights_comparison": rights_summary,
                    "risk_comparison": {
                        "contract_a_risk": risk_comp.get('contract_a_risk_level', 'N/A'),
                        "contract_b_risk": risk_comp.get('contract_b_risk_level', 'N/A'),
                        "lower_risk": risk_comp.get('lower_label_risk', 'N/A')
                    },
                    "key_differences": key_diffs,
                    "deal_breakers": deal_breakers if deal_breakers else "None identified",
                    "overall_assessment": result.get('overall_assessment', ''),
                    "recommended_terms": result.get('recommended_standard_terms', [])[:3]
                }

                return json.dumps({
                    "status": "success",
                    "action": "compare_contracts",
                    "contract_a": contract_name,
                    "contract_b": contract_name_b,
                    "chat_response": chat_summary,
                    "_full_data": {"comparison": result}
                }, indent=2)

            else:
                return json.dumps({
                    "status": "error",
                    "message": f"Unknown action: {action}",
                    "valid_actions": ["list_contracts", "analyze_contract", "extract_clauses",
                                     "summarize_contract", "identify_risks", "compare_contracts"]
                }, indent=2)

        except Exception as e:
            logging.error(f"ContractAnalysisAgent error: {e}")
            return json.dumps({
                "status": "error",
                "message": str(e),
                "type": type(e).__name__
            }, indent=2)


if __name__ == "__main__":
    # Test the agent
    agent = ContractAnalysisAgent()

    print("Testing ContractAnalysisAgent...")
    print("\n1. Listing contracts:")
    print(agent.perform(action="list_contracts"))
