import os
import json
import logging
import urllib.parse
import re
import base64
from datetime import datetime
import requests
import msal
from io import BytesIO
from agents.basic_agent import BasicAgent
from utils.azure_file_storage import AzureFileStorageManager
from openai import AzureOpenAI

# Optional imports - will be attempted at runtime
DOCX_AVAILABLE = False
PDF_AVAILABLE = False
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    try:
        from pypdf import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        pass

# Try to import image processing libraries
IMAGE_PROCESSING_AVAILABLE = False
try:
    from PIL import Image
    import io
    IMAGE_PROCESSING_AVAILABLE = True
except ImportError:
    pass

class SharePointDocumentExtractorAgent(BasicAgent):
    def __init__(self):
        self.name = "SharePointDocumentExtractor"
        self.metadata = {
            "name": self.name,
            "description": "Extracts complete content from SharePoint documents using direct URLs. Supports various file types including DOCX, PDF, and images. All extracted content is stored in Azure File Storage for future reference.",
            "parameters": {
                "type": "object",
                "properties": {
                    "document_url": {
                        "type": "string",
                        "description": "The full URL of the SharePoint document or folder (e.g., https://tenant.sharepoint.com/sites/site/Shared%20Documents/folder/document.docx). If not provided or set to 'default', uses environment variables."
                    },
                    "analyze_images": {
                        "type": "boolean",
                        "description": "Whether to analyze images found in the document or folder using Vision AI",
                        "default": False
                    },
                    "extract_full_content": {
                        "type": "boolean",
                        "description": "Whether to extract and return the full content instead of just a preview",
                        "default": True
                    }
                },
                "required": []
            }
        }
        super().__init__(name=self.name, metadata=self.metadata)

        # Setup logging
        self.logger = logging.getLogger(self.name)

        # Load Azure AD app credentials from environment variables
        self.client_id = os.environ.get('SHAREPOINT_CLIENT_ID')
        self.client_secret = os.environ.get('SHAREPOINT_CLIENT_SECRET')
        self.tenant_id = os.environ.get('SHAREPOINT_TENANT_ID')

        # Load SharePoint configuration from environment variables
        self.default_tenant_url = os.environ.get('SHAREPOINT_TENANT_URL', '')
        self.default_site_name = os.environ.get('SHAREPOINT_SITE_NAME', '')
        self.default_document_path = os.environ.get('SHAREPOINT_DOCUMENT_PATH', 'Shared%20Documents')

        # Token for authentication
        self.access_token = None

        # Initialize Azure File Storage Manager
        try:
            self.storage_manager = AzureFileStorageManager()
        except Exception as e:
            self.logger.error(f"Error initializing AzureFileStorageManager: {str(e)}")
            # Create a direct connection to Azure File Storage as fallback
            storage_connection = os.environ.get('AzureWebJobsStorage', '')
            if storage_connection:
                try:
                    from azure.storage.file import FileService
                    connection_parts = dict(part.split('=', 1) for part in storage_connection.split(';'))
                    self.account_name = connection_parts.get('AccountName')
                    self.account_key = connection_parts.get('AccountKey')
                    self.share_name = os.environ.get('AZURE_FILES_SHARE_NAME', 'azfbusinessbot')
                    self.file_service = FileService(
                        account_name=self.account_name,
                        account_key=self.account_key
                    )
                    self.logger.info("Successfully created direct Azure File Storage connection as fallback")
                except Exception as e2:
                    self.logger.error(f"Failed to create direct Azure File Storage connection: {str(e2)}")

    def _parse_sharepoint_url(self, document_url):
        """
        Parse a SharePoint document URL into its components.
        Returns a tuple of (tenant_url, site_name, document_path, file_name, is_folder)
        """
        # Use environment variables if document_url is not provided or is 'default'
        if not document_url or document_url.lower() == "default":
            if not self.default_tenant_url or not self.default_site_name:
                raise ValueError("SharePoint configuration not found in environment variables. Please set SHAREPOINT_TENANT_URL and SHAREPOINT_SITE_NAME.")

            tenant_url = self.default_tenant_url
            site_name = self.default_site_name
            document_path = self.default_document_path

            # Construct the full URL for logging
            document_url = f"{tenant_url}/sites/{site_name}/{document_path}"
            self.logger.info(f"Using default SharePoint configuration: {document_url}")

            # For default configuration, assume it's a folder
            file_name = ""
            is_folder = True

            return (tenant_url, site_name, document_path, file_name, is_folder)

        try:
            # Parse the provided URL
            parsed_url = urllib.parse.urlparse(document_url)

            # Extract the hostname (tenant)
            tenant_url = f"{parsed_url.scheme}://{parsed_url.netloc}"

            # URL decode the path
            server_relative_url = urllib.parse.unquote(parsed_url.path)

            # Extract the site name
            site_match = re.search(r'/sites/([^/]+)', server_relative_url)
            site_name = site_match.group(1) if site_match else None

            # Extract the document path (relative to the site)
            if site_name:
                site_path = f"/sites/{site_name}"
                document_path = server_relative_url.replace(site_path, "").lstrip("/")
            else:
                document_path = server_relative_url.lstrip("/")

            # Extract the file name
            file_name = os.path.basename(document_path)

            # Check if the URL points to a folder (no file extension)
            is_folder = '.' not in file_name or file_name.endswith('/')

            return (tenant_url, site_name, document_path, file_name, is_folder)
        except Exception as e:
            self.logger.error(f"Error parsing SharePoint URL: {str(e)}")
            raise ValueError(f"Invalid SharePoint document URL: {document_url}")

    def authenticate(self):
        """
        Authenticate using Microsoft Entra ID (Azure AD)
        """
        try:
            self.logger.info(f"Authenticating with Entra ID using client ID: {self.client_id[:5]}... and tenant ID: {self.tenant_id[:8]}...")
            # Create a confidential client application
            app = msal.ConfidentialClientApplication(
                client_id=self.client_id,
                client_credential=self.client_secret,
                authority=f"https://login.microsoftonline.com/{self.tenant_id}"
            )

            # Acquire token for Microsoft Graph
            result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])

            if "access_token" in result:
                self.logger.info("Successfully acquired token for Microsoft Graph")
                self.access_token = result["access_token"]
                return self.access_token
            else:
                self.logger.error(f"Failed to acquire token: {result.get('error_description', 'Unknown error')}")
                raise Exception(f"Failed to authenticate: {result.get('error_description', 'Unknown error')}")
        except Exception as e:
            self.logger.error(f"Authentication error: {str(e)}")
            raise

    def _get_file_content(self, tenant_url, site_name, document_path, file_name):
        """
        Get file content using Microsoft Graph API
        """
        try:
            # Make sure we have a valid token
            if not self.access_token:
                self.authenticate()

            # Extract the tenant hostname
            tenant_hostname = tenant_url.split('//')[1]

            # Set up headers with the access token
            headers = {
                'Authorization': f'Bearer {self.access_token}',
                'Accept': '*/*'
            }

            # First, let's get the site ID using the site address
            site_lookup_url = f"https://graph.microsoft.com/v1.0/sites/{tenant_hostname}:/sites/{site_name}"
            self.logger.info(f"Looking up site ID: {site_lookup_url}")

            site_response = requests.get(site_lookup_url, headers=headers)

            if site_response.status_code == 200:
                site_data = site_response.json()
                site_id = site_data.get('id')
                self.logger.info(f"Found site ID: {site_id}")

                # Now try to get the default document library (usually "documents")
                drives_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives"
                self.logger.info(f"Looking up drives: {drives_url}")

                drives_response = requests.get(drives_url, headers=headers)

                if drives_response.status_code == 200:
                    drives_data = drives_response.json()
                    if 'value' in drives_data and len(drives_data['value']) > 0:
                        # Use the first drive (usually Documents)
                        drive_id = drives_data['value'][0]['id']
                        self.logger.info(f"Found drive ID: {drive_id}")

                        # Normalize path - remove leading/trailing slashes and spaces
                        clean_path = document_path.strip('/ ')

                        # Special handling for "Shared Documents" which often needs to be reformatted
                        if clean_path.startswith("Shared Documents/"):
                            # Try without "Shared Documents" prefix as it's often implicitly the root
                            file_without_shared = clean_path.replace("Shared Documents/", "")
                            file_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{file_without_shared}:/content"
                        else:
                            file_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{clean_path}:/content"

                        self.logger.info(f"Trying to get file: {file_url}")

                        file_response = requests.get(file_url, headers=headers)

                        if file_response.status_code == 200:
                            self.logger.info(f"Successfully retrieved file content: {len(file_response.content)} bytes")
                            return file_response.content

                        # If we failed to get the file directly, try alternative approaches
                        self.logger.warning(f"Failed to get file. Status: {file_response.status_code}. Trying alternatives...")

                        # Try looking for the file by name in the root directory
                        list_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root/children"
                        self.logger.info(f"Listing root files: {list_url}")

                        list_response = requests.get(list_url, headers=headers)
                        if list_response.status_code == 200:
                            list_data = list_response.json()

                            # Try to find our file in the listing
                            for item in list_data.get('value', []):
                                if item.get('name') == file_name:
                                    self.logger.info(f"Found file in listing! ID: {item.get('id')}")

                                    # Get file content by ID
                                    item_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/items/{item.get('id')}/content"
                                    self.logger.info(f"Getting file by ID: {item_url}")

                                    item_response = requests.get(item_url, headers=headers)
                                    if item_response.status_code == 200:
                                        self.logger.info(f"Successfully retrieved file by ID: {len(item_response.content)} bytes")
                                        return item_response.content

            # If all the approaches above failed, try additional patterns as a last resort
            patterns = [
                f"https://graph.microsoft.com/v1.0/sites/{tenant_hostname}:/sites/{site_name}:/drive/root:/{document_path}:/content",
                f"https://graph.microsoft.com/v1.0/sites/{tenant_hostname}:/sites/{site_name}:/drives/documentLibrary/root:/{document_path}:/content",
                f"https://graph.microsoft.com/v1.0/sites/{tenant_hostname}:/sites/{site_name}:/drives/documents/root:/{document_path}:/content"
            ]

            for pattern in patterns:
                self.logger.info(f"Trying pattern: {pattern}")
                response = requests.get(pattern, headers=headers)
                if response.status_code == 200:
                    self.logger.info(f"Pattern worked! Retrieved {len(response.content)} bytes")
                    return response.content

            self.logger.error("All patterns failed to retrieve file content")
            return None
        except Exception as e:
            self.logger.error(f"Error in _get_file_content: {str(e)}")
            return None

    def _list_folder_contents(self, tenant_url, site_name, folder_path):
        """
        Get a list of all files in a SharePoint folder
        """
        try:
            # Make sure we have a valid token
            if not self.access_token:
                self.authenticate()

            # Extract the tenant hostname
            tenant_hostname = tenant_url.split('//')[1]

            # Set up headers with the access token
            headers = {
                'Authorization': f'Bearer {self.access_token}',
                'Accept': 'application/json'
            }

            # First, let's get the site ID using the site address
            site_lookup_url = f"https://graph.microsoft.com/v1.0/sites/{tenant_hostname}:/sites/{site_name}"
            self.logger.info(f"Looking up site ID for folder listing: {site_lookup_url}")

            site_response = requests.get(site_lookup_url, headers=headers)

            if site_response.status_code == 200:
                site_data = site_response.json()
                site_id = site_data.get('id')
                self.logger.info(f"Found site ID: {site_id}")

                # Get the drives (document libraries)
                drives_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives"
                drives_response = requests.get(drives_url, headers=headers)

                if drives_response.status_code == 200:
                    drives_data = drives_response.json()
                    if 'value' in drives_data and len(drives_data['value']) > 0:
                        # Use the first drive (usually Documents)
                        drive_id = drives_data['value'][0]['id']
                        self.logger.info(f"Found drive ID: {drive_id}")

                        # For Shared Documents, just use the root of the drive
                        # This is a special case for the default document library
                        if folder_path.strip('/') == 'Shared Documents' or folder_path.strip('/') == 'Shared Documents/':
                            folder_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root/children"
                            self.logger.info(f"Listing root children (Shared Documents): {folder_url}")
                        else:
                            # For other paths
                            clean_path = folder_path.strip('/ ')

                            # Special handling for "Shared Documents" prefix
                            if clean_path.startswith("Shared Documents/"):
                                # Remove the "Shared Documents" prefix
                                folder_without_shared = clean_path.replace("Shared Documents/", "")
                                if folder_without_shared:
                                    folder_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{folder_without_shared}:/children"
                                else:
                                    # If we're looking at the root of Shared Documents
                                    folder_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root/children"
                            elif clean_path:
                                folder_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{clean_path}:/children"
                            else:
                                folder_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root/children"

                        self.logger.info(f"Listing folder contents: {folder_url}")

                        folder_response = requests.get(folder_url, headers=headers)

                        if folder_response.status_code == 200:
                            folder_data = folder_response.json()
                            items = folder_data.get('value', [])
                            self.logger.info(f"Found {len(items)} items in folder")
                            return items
                        else:
                            self.logger.error(f"Failed to list folder. Status: {folder_response.status_code}, Response: {folder_response.text}")

                            # Try a direct approach as a fallback
                            self.logger.info("Trying fallback approach to list files...")
                            try:
                                fallback_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/items/root/children"
                                self.logger.info(f"Fallback URL: {fallback_url}")
                                fallback_response = requests.get(fallback_url, headers=headers)

                                if fallback_response.status_code == 200:
                                    fallback_data = fallback_response.json()
                                    items = fallback_data.get('value', [])
                                    self.logger.info(f"Fallback approach found {len(items)} items")
                                    return items
                            except Exception as e:
                                self.logger.error(f"Error in fallback approach: {str(e)}")

            return []
        except Exception as e:
            self.logger.error(f"Error in _list_folder_contents: {str(e)}")
            return []

    def _extract_text(self, file_content, file_ext, file_name):
        """Extract text from different file types"""
        try:
            if file_ext.lower() == '.txt':
                # For text files, just decode the content
                return file_content.decode('utf-8', errors='replace')

            elif file_ext.lower() == '.docx':
                self.logger.info(f"Extracting text from DOCX file: {file_name}")

                # First check if python-docx is available
                if not DOCX_AVAILABLE:
                    self.logger.warning("python-docx package is not installed, returning raw content as text")
                    # Return something rather than failing completely
                    return f"[Document content from {file_name}] - Raw content available but python-docx package is not installed for proper extraction"

                # For Word documents
                try:
                    file_stream = BytesIO(file_content)
                    doc = Document(file_stream)

                    # Extract paragraphs
                    paragraphs = [para.text for para in doc.paragraphs if para.text]

                    # Extract tables
                    tables_text = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text for cell in row.cells]
                            tables_text.append(" | ".join(row_text))

                    # Combine all text
                    all_text = paragraphs + tables_text
                    result = "\n".join(all_text)

                    self.logger.info(f"Successfully extracted {len(result)} characters from DOCX")

                    if not result.strip():
                        self.logger.warning("Extracted text is empty, document might be empty or contain only images")
                        return f"[Document appears to be empty or contains only images/non-text content: {file_name}]"

                    return result
                except Exception as e:
                    self.logger.error(f"Error with python-docx: {str(e)}")
                    # Fall back to raw content
                    return f"[Document content from {file_name}] - Error extracting with python-docx: {str(e)}"

            elif file_ext.lower() == '.pdf':
                self.logger.info(f"Extracting text from PDF file: {file_name}")

                if not PDF_AVAILABLE:
                    self.logger.warning("PDF processing package is not installed, returning raw content as text")
                    return f"[Document content from {file_name}] - Raw content available but PDF package is not installed for proper extraction"

                # For PDF files
                file_stream = BytesIO(file_content)

                try:
                    # Try PyPDF2 first
                    pdf = pypdf.PdfReader(file_stream)
                    text = []
                    for page_num in range(len(pdf.pages)):
                        page = pdf.pages[page_num]
                        text.append(page.extract_text())

                    result = "\n".join(text)
                    self.logger.info(f"Successfully extracted {len(result)} characters from PDF")

                    if not result.strip():
                        self.logger.warning("Extracted text is empty, PDF might be scanned or contain only images")
                        return f"[PDF may be scanned or contain only images: {file_name}]"

                    return result
                except NameError:
                    # Try pypdf if PyPDF2 is not available
                    try:
                        pdf = PdfReader(file_stream)
                        text = []
                        for page in pdf.pages:
                            text.append(page.extract_text())

                        result = "\n".join(text)
                        self.logger.info(f"Successfully extracted {len(result)} characters from PDF")

                        if not result.strip():
                            self.logger.warning("Extracted text is empty, PDF might be scanned or contain only images")
                            return f"[PDF may be scanned or contain only images: {file_name}]"

                        return result
                    except Exception as e:
                        self.logger.error(f"Error with pypdf: {str(e)}")
                        return f"[Document content from {file_name}] - Error extracting with pypdf: {str(e)}"

            elif file_ext.lower() in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                # For image files, return a placeholder - actual analysis will be done separately
                return f"[Image file: {file_name}]"

            else:
                self.logger.warning(f"Unsupported file type: {file_ext}")
                return f"[Document content from {file_name}] - File type {file_ext} is not supported for text extraction"

        except Exception as e:
            self.logger.error(f"Error extracting text from '{file_name}': {str(e)}")
            # Return something meaningful rather than just failing
            return f"[Document content from {file_name}] - Error during extraction: {str(e)}"

    def _analyze_image(self, image_content, image_name):
        """Analyze an image using Vision AI"""
        try:
            # Base64 encode the image
            encoded_image = base64.b64encode(image_content).decode('utf-8')

            # Initialize the Azure OpenAI client
            client = AzureOpenAI(
                api_key=os.environ['AZURE_OPENAI_API_KEY'],
                api_version=os.environ['AZURE_OPENAI_API_VERSION'],
                azure_endpoint=os.environ['AZURE_OPENAI_ENDPOINT']
            )

            # Create the messages payload for the API request
            messages = [
                {
                    "role": "system",
                    "content": "You are an AI assistant that analyzes images and provides detailed descriptions."
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{encoded_image}"
                            }
                        },
                        {
                            "type": "text",
                            "text": f"Analyze this image named '{image_name}' and provide a detailed description."
                        }
                    ]
                }
            ]

            # Request analysis from the model
            response = client.chat.completions.create(
                model="gpt-4o-mini",  # Use vision-enabled model
                messages=messages,
                temperature=0.7,
                max_tokens=800
            )

            analysis = response.choices[0].message.content

            self.logger.info(f"Successfully analyzed image {image_name}")
            return analysis
        except Exception as e:
            self.logger.error(f"Error analyzing image '{image_name}': {str(e)}")
            return f"Error analyzing image: {str(e)}"

    def _store_content_in_azure_files(self, site_url, file_name, extracted_text, image_analysis=None):
        """Store the extracted content in Azure File Storage"""
        try:
            # Create a base path based on the SharePoint URL
            parsed_url = urllib.parse.urlparse(site_url)
            domain = parsed_url.netloc

            # Generate a timestamp for this extraction
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # Create a safe storage path
            safe_file_name = file_name.replace('/', '_').replace('\\', '_')
            storage_path = f"sharepoint_docs/{domain}/{timestamp}"

            # Check if we're using the AzureFileStorageManager or direct FileService
            if hasattr(self, 'storage_manager'):
                # Use the storage manager
                self.storage_manager.ensure_directory_exists(storage_path)

                # Store the summary JSON file
                summary_data = {
                    "extraction_time": timestamp,
                    "sharepoint_site": site_url,
                    "file_name": file_name,
                    "has_image_analysis": image_analysis is not None
                }

                self.storage_manager.write_file(
                    storage_path,
                    "summary.json",
                    json.dumps(summary_data, indent=2)
                )

                # Store the file content
                self.storage_manager.write_file(
                    storage_path,
                    f"{safe_file_name}.txt",
                    extracted_text
                )

                # Store image analysis if available
                if image_analysis:
                    self.storage_manager.write_file(
                        storage_path,
                        f"{safe_file_name}_image_analysis.json",
                        json.dumps(image_analysis, indent=2)
                    )

                content_path = f"{storage_path}/{safe_file_name}.txt"

            elif hasattr(self, 'file_service'):
                # Use direct FileService
                # Make sure the share exists
                self.file_service.create_share(self.share_name, fail_on_exist=False)

                # Create the directory path
                path_parts = storage_path.split('/')
                current_path = ""
                for part in path_parts:
                    if part:
                        if current_path:
                            current_path = f"{current_path}/{part}"
                        else:
                            current_path = part
                        self.file_service.create_directory(
                            self.share_name,
                            current_path,
                            fail_on_exist=False
                        )

                # Store the summary JSON file
                summary_data = {
                    "extraction_time": timestamp,
                    "sharepoint_site": site_url,
                    "file_name": file_name,
                    "has_image_analysis": image_analysis is not None
                }

                self.file_service.create_file_from_text(
                    self.share_name,
                    storage_path,
                    "summary.json",
                    json.dumps(summary_data, indent=2)
                )

                # Store the content
                self.file_service.create_file_from_text(
                    self.share_name,
                    storage_path,
                    f"{safe_file_name}.txt",
                    extracted_text
                )

                # Store image analysis if available
                if image_analysis:
                    self.file_service.create_file_from_text(
                        self.share_name,
                        storage_path,
                        f"{safe_file_name}_image_analysis.json",
                        json.dumps(image_analysis, indent=2)
                    )

                content_path = f"{storage_path}/{safe_file_name}.txt"
            else:
                return {
                    "status": "error",
                    "message": "No Azure File Storage connection available"
                }

            return {
                "status": "success",
                "storage_path": storage_path,
                "content_path": content_path
            }

        except Exception as e:
            self.logger.error(f"Error storing content in Azure File Storage: {str(e)}")
            return {
                "status": "error",
                "message": f"Failed to store content in Azure File Storage: {str(e)}"
            }

    def perform(self, **kwargs):
        try:
            # Get document URL from kwargs, or use 'default' to trigger environment variable usage
            document_url = kwargs.get('document_url', 'default')

            # Get other parameters from kwargs
            analyze_images = kwargs.get('analyze_images', False)
            extract_full_content = kwargs.get('extract_full_content', True)

            # Validate required environment variables
            if not self.client_id or not self.client_secret or not self.tenant_id:
                return json.dumps({
                    "status": "error",
                    "message": "Entra ID credentials not found. Please set SHAREPOINT_CLIENT_ID, SHAREPOINT_CLIENT_SECRET and SHAREPOINT_TENANT_ID environment variables."
                })

            # Parse the document URL
            try:
                tenant_url, site_name, document_path, file_name, is_folder = self._parse_sharepoint_url(document_url)
                self.logger.info(f"SharePoint URL parsed - Tenant: {tenant_url}, Site: {site_name}, Path: {document_path}, File: {file_name}, Is Folder: {is_folder}")
            except ValueError as e:
                return json.dumps({
                    "status": "error",
                    "message": str(e)
                })

            # Authenticate with Microsoft Entra ID
            try:
                self.authenticate()
            except Exception as e:
                error_msg = str(e)
                self.logger.error(f"Failed to authenticate with Entra ID: {error_msg}")
                return json.dumps({
                    "status": "error",
                    "message": f"Failed to authenticate with Entra ID: {error_msg}"
                })

            # Handle different extractions based on whether it's a folder or single file
            if is_folder:
                # Get all items in the folder
                folder_items = self._list_folder_contents(tenant_url, site_name, document_path)

                if not folder_items:
                    return json.dumps({
                        "status": "error",
                        "message": f"Failed to list items in folder: {document_path} or the folder is empty"
                    })

                # Process each item in the folder
                results = []
                all_extracted_text = ""
                image_analysis_results = {}

                for item in folder_items:
                    item_name = item.get('name', '')
                    item_id = item.get('id', '')
                    item_type = 'folder' if item.get('folder') else 'file'

                    # Skip folders for now
                    if item_type == 'folder':
                        continue

                    # Get file content
                    try:
                        # Get the file content directly using the item ID if available
                        item_url = f"https://graph.microsoft.com/v1.0/drives/{item.get('parentReference', {}).get('driveId')}/items/{item_id}/content"
                        headers = {
                            'Authorization': f'Bearer {self.access_token}',
                            'Accept': '*/*'
                        }

                        item_response = requests.get(item_url, headers=headers)

                        if item_response.status_code == 200:
                            file_content = item_response.content
                            self.logger.info(f"Successfully retrieved {item_name}: {len(file_content)} bytes")

                            # Determine file extension
                            file_ext = os.path.splitext(item_name)[1]

                            # Extract text from the file
                            extracted_text = self._extract_text(file_content, file_ext, item_name)

                            # Add to the combined text with a header for each file
                            all_extracted_text += f"\n\n### FILE: {item_name} ###\n\n"
                            all_extracted_text += extracted_text

                            # Analyze image if requested and it's an image file
                            image_analysis = None
                            if analyze_images and file_ext.lower() in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                                self.logger.info(f"Analyzing image: {item_name}")
                                image_analysis = self._analyze_image(file_content, item_name)
                                image_analysis_results[item_name] = image_analysis
                                all_extracted_text += f"\n\n### IMAGE ANALYSIS: {item_name} ###\n\n"
                                all_extracted_text += image_analysis

                            # Create a result entry with either full content or preview
                            if extract_full_content:
                                content_for_result = extracted_text
                            else:
                                content_for_result = extracted_text[:1000] + ("..." if len(extracted_text) > 1000 else "")

                            results.append({
                                "item_name": item_name,
                                "item_type": item_type,
                                "content": content_for_result,
                                "has_image_analysis": image_analysis is not None
                            })
                        else:
                            self.logger.error(f"Failed to get item content. Status: {item_response.status_code}")
                            all_extracted_text += f"\n\n### ERROR: {item_name} ###\n\n"
                            all_extracted_text += f"Failed to retrieve content. Status code: {item_response.status_code}"
                    except Exception as e:
                        self.logger.error(f"Error processing item {item_name}: {str(e)}")
                        all_extracted_text += f"\n\n### ERROR: {item_name} ###\n\n"
                        all_extracted_text += f"Error: {str(e)}"
                        results.append({
                            "item_name": item_name,
                            "item_type": item_type,
                            "error": str(e)
                        })

                # Combine all extracted text and image analysis for storage
                combined_text = "# Folder Contents Extraction\n\n"
                combined_text += f"Folder: {document_path}\n\n"
                combined_text += f"Number of items: {len(results)}\n\n"
                combined_text += f"Extraction time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
                combined_text += all_extracted_text

                # Store the results
                storage_result = self._store_content_in_azure_files(
                    tenant_url,
                    os.path.basename(document_path) or "folder",
                    combined_text,
                    image_analysis_results if image_analysis_results else None
                )

                # Return results with either full content or preview
                if extract_full_content:
                    # If full content is requested, include it in the response
                    return json.dumps({
                        "status": "success",
                        "message": f"Successfully extracted content from {len(results)} items in folder",
                        "folder_info": {
                            "folder_path": document_path,
                            "items_processed": len(results)
                        },
                        "storage": storage_result,
                        "full_content": combined_text
                    })
                else:
                    # Otherwise just return the summary with item previews
                    return json.dumps({
                        "status": "success",
                        "message": f"Successfully extracted content from {len(results)} items in folder",
                        "folder_info": {
                            "folder_path": document_path,
                            "items_processed": len(results)
                        },
                        "items": results,
                        "storage": storage_result
                    })
            else:
                # Single file extraction
                # Get the file content using Microsoft Graph API
                try:
                    file_content = self._get_file_content(tenant_url, site_name, document_path, file_name)
                    if not file_content:
                        return json.dumps({
                            "status": "error",
                            "message": f"Failed to retrieve file content for {file_name}"
                        })
                    self.logger.info(f"Retrieved file content: {len(file_content)} bytes")
                except Exception as e:
                    error_msg = str(e)
                    self.logger.error(f"Failed to get file content: {error_msg}")
                    return json.dumps({
                        "status": "error",
                        "message": f"Failed to get file content: {error_msg}"
                    })

                # Extract text based on file type
                file_ext = os.path.splitext(file_name)[1].lower()
                extracted_text = self._extract_text(file_content, file_ext, file_name)

                # Analyze image if requested and it's an image file
                image_analysis = None
                complete_content = extracted_text

                if analyze_images and file_ext.lower() in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                    self.logger.info(f"Analyzing image: {file_name}")
                    image_analysis = self._analyze_image(file_content, file_name)
                    # Add image analysis to the complete content
                    complete_content += f"\n\n### IMAGE ANALYSIS ###\n\n{image_analysis}"

                if not extracted_text and not image_analysis:
                    return json.dumps({
                        "status": "error",
                        "message": f"Failed to extract content from file: {file_name}"
                    })

                # Store the extracted content in Azure File Storage
                storage_result = self._store_content_in_azure_files(
                    tenant_url,
                    file_name,
                    complete_content,
                    {"image_analysis": image_analysis} if image_analysis else None
                )

                # Build response
                response = {
                    "status": "success",
                    "message": f"Successfully extracted content from {file_name}",
                    "file_info": {
                        "file_name": file_name,
                        "file_url": document_url if document_url != 'default' else f"{tenant_url}/sites/{site_name}/{document_path}",
                        "file_size_bytes": len(file_content)
                    },
                    "storage": storage_result
                }

                # Include full content or just a preview based on parameter
                if extract_full_content:
                    response["full_content"] = complete_content
                else:
                    response["content_preview"] = extracted_text[:5000] + ("..." if len(extracted_text) > 5000 else "")
                    if image_analysis:
                        response["image_analysis_preview"] = image_analysis[:500] + ("..." if len(image_analysis) > 500 else "")

                return json.dumps(response)

        except Exception as e:
            error_msg = str(e)
            self.logger.error(f"Unhandled error in agent execution: {error_msg}")
            return json.dumps({
                "status": "error",
                "message": f"Unhandled error: {error_msg}"
            })
